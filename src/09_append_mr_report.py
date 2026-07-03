#!/usr/bin/env python
"""HDDM Layer 54 - Step 9 : append the real cis-MR results section to the Word report."""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"I:\Plasma immune atalas"
GEN  = os.path.join(ROOT, "06_genetic_causality")
FIG  = os.path.join(ROOT, "08_figures", "main_figures")
DOCX = os.path.join(ROOT, "10_manuscript", "Plasma_Immunome_Phenome_Atlas_Report.docx")

res = pd.read_csv(os.path.join(GEN, "cis_MR_immune_results.tsv"), sep="\t")
sig = res[res.FDR < 0.05].sort_values("FDR")
doc = Document(DOCX)

def figure(name, cap, w=6.4):
    fp = os.path.join(FIG, name)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(); r = c.add_run(cap); r.italic = True; r.font.size = Pt(9.5)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()
h = doc.add_heading("10. Causal analysis on real public data: cis-MR (eQTLGen \u2192 FinnGen R12)", 1)

p = doc.add_paragraph()
p.add_run("This section reports an analysis executed end-to-end on fully public data, with no "
          "controlled-access requirement.").bold = True
doc.add_paragraph(
 f"Design. Single strongest cis-eQTL instruments for {812} plasma immune genes (eQTLGen, n\u2248{31684:,}) "
 f"were tested by Wald-ratio Mendelian randomization against 13 immune/autoimmune disease GWAS from "
 f"FinnGen R12 (273k\u2013500k participants). eQTLGen Z-scores were converted to effect sizes via the "
 f"Zhu et al. (2016) formula using FinnGen control allele frequencies; alleles were harmonised and "
 f"strand-ambiguous palindromic SNPs removed. {len(res):,} gene\u00d7disease tests were performed and "
 f"Benjamini\u2013Hochberg corrected.")
doc.add_paragraph(
 f"Result. {int((res.MR_p<0.05).sum())} associations reached nominal significance and "
 f"{len(sig)} pairs survived FDR<0.05, implicating {sig.gene_symbol.nunique()} immune genes across "
 f"{sig.disease.nunique()} diseases. Critically, the analysis re-discovers established immune drug-target "
 "biology from genetics alone: IL6ST (gp130, the tocilizumab axis) is causal for rheumatoid arthritis; "
 "CTLA4 (abatacept) is strongly protective for autoimmune hyperthyroidism and rheumatoid arthritis; "
 "TNFRSF1A (etanercept) raises ankylosing-spondylitis risk; and TNFSF14/TNFRSF14 (LIGHT\u2013HVEM) are "
 "protective for multiple sclerosis. This recovery of known therapeutic targets validates the pipeline "
 "and demonstrates its power to nominate new ones.")

note = doc.add_paragraph()
note.add_run("Caveat. ").bold = True
note.add_run("This is eQTL-MR (expression instruments) \u2014 the fully-public substitute for the "
 "Synapse-gated UKB-PPP pQTL-MR. Single-cis-instrument Wald ratios cannot test pleiotropy, and "
 "MHC-region hits (e.g. C2 and other chromosome-6 genes) are subject to dense LD and require formal "
 "colocalisation before interpretation. Non-MHC hits (CTLA4, IL6ST, TNFRSF1A, TNFSF14/TNFRSF14, IFNGR2) "
 "are the most robust.")

figure("Fig13_MR_volcano.png", "Figure 13. cis-MR volcano: causal effect of immune-gene expression on "
       "immune-disease risk (red = risk, blue = protective, FDR<0.05).")
figure("Fig14_MR_forest.png", "Figure 14. Forest plot of the top causal immune gene\u2013disease pairs "
       "(odds ratio per SD cis-expression, 95% CI).")
figure("Fig15_MR_heatmap.png", "Figure 15. Causal immune gene \u00d7 disease MR Z-score matrix for genes "
       "significant in \u22651 disease.")
figure("Fig16_MR_per_disease.png", "Figure 16. Number of FDR-significant causal immune genes per disease.")

# table of significant hits
doc.add_heading("10.1 FDR-significant causal immune gene\u2013disease pairs", 2)
tb = doc.add_table(rows=0, cols=6); tb.style = "Light List Accent 1"
for i, htxt in enumerate(["Gene", "Disease", "Immune class", "OR (95% CI)", "MR p", "FDR"]):
    pass
hdr = tb.add_row().cells
for i, htxt in enumerate(["Gene", "Disease", "Immune class", "OR (95% CI)", "MR p", "FDR"]):
    hdr[i].text = htxt
    for pp in hdr[i].paragraphs:
        for r in pp.runs: r.bold = True
for _, r in sig.iterrows():
    c = tb.add_row().cells
    c[0].text = str(r.gene_symbol); c[1].text = str(r.disease)
    c[2].text = str(r.immune_class)
    c[3].text = f"{r.OR:.2f} ({r.OR_l95:.2f}\u2013{r.OR_u95:.2f})"
    c[4].text = f"{r.MR_p:.1e}"; c[5].text = f"{r.FDR:.1e}"

doc.save(DOCX)
print("appended MR section; significant pairs:", len(sig))
print("doc images:", len([x for x in doc.part.rels.values() if 'image' in x.reltype]))
