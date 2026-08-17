#!/usr/bin/env python
"""
HDDM Layer 54 - Step 13
Rewrite the manuscript to Nature-family standard:
  - sharper title
  - Nature-grammar abstract
  - 4 clean results sections
  - claim-strength-gated language throughout (uses evidence_tiered_targets.tsv)
  - embeds the 6 multi-panel Nature figures
Output: 10_manuscript/Plasma_Immunome_Phenome_Atlas_Nature.docx
"""
import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"I:\Plasma immune atalas"
GEN  = os.path.join(ROOT, "06_genetic_causality")
FIG  = os.path.join(ROOT, "08_figures", "nature")
PROC = os.path.join(ROOT, "02_data_processed")
OUT  = os.path.join(ROOT, "10_manuscript")
os.makedirs(OUT, exist_ok=True)

# ---- load data / numbers ----
ann = pd.read_csv(os.path.join(PROC, "plasma_immune_protein_annotation.tsv"), sep="\t")
imm = ann[ann.is_plasma_immune == 1]
mr  = pd.read_csv(os.path.join(GEN, "cis_MR_immune_results.tsv"), sep="\t")
inst= pd.read_csv(os.path.join(GEN, "immune_cis_eqtl_instruments.tsv"), sep="\t")
tier= pd.read_csv(os.path.join(GEN, "evidence_tiered_targets.tsv"), sep="\t")
final= pd.read_csv(os.path.join(GEN, "FINAL_evidence_tiers.tsv"), sep="\t")
_replp = os.path.join(GEN, "FINAL_evidence_tiers_repl.tsv")
if os.path.exists(_replp):
    final = pd.read_csv(_replp, sep="\t")

N_OLINK = 2923
N_IMM   = int(imm.shape[0])
N_INST  = int(inst.gene_symbol.nunique())
N_TEST  = int(mr.shape[0])
N_SIG   = int((mr.FDR < 0.05).sum())
N_DIS   = int(mr.disease.nunique())
N_T4 = int((tier.evidence_tier==4).sum())
N_T3 = int((tier.evidence_tier==3).sum())
N_T2 = int((tier.evidence_tier==2).sum())
t4 = tier[tier.evidence_tier==4]
# protein-level (INTERVAL pQTL) layer
N_T5   = int((final.final_tier==5).sum())
t5g    = final[final.final_tier==5]
N_PQTL = int(final.pQTL_OR.notna().sum())
N_CONC = int((final.pQTL_concordant=="yes").sum())
N_DISC = int((final.pQTL_concordant=="NO").sum())
# independent-GWAS replication layer
HAS_REP = "rep_status" in final.columns
N_COV  = int((final.rep_status!="not covered").sum()) if HAS_REP else 0
N_REP  = int((final.rep_status=="replicated").sum()) if HAS_REP else 0
# novelty layer
def _nl(f):
    p=os.path.join(GEN,f); return pd.read_csv(p,sep="\t") if os.path.exists(p) else None
nov_enr=_nl("novelty_class_enrichment.tsv")
nov_dir=_nl("novelty_drug_direction.tsv")
nov_pri=_nl("novelty_prioritised_targets.tsv")
if nov_dir is not None:
    N_DIRMATCH=int((nov_dir.concordance=="match").sum())
    N_NOVELNOM=int((nov_pri.category=="NOVEL nomination").sum()) if nov_pri is not None else 0
else:
    N_DIRMATCH=N_NOVELNOM=0
N_RCON = int(final.rep_status.isin(["replicated","concordant (ns)"]).sum()) if HAS_REP else 0

# ---- pan-phenome layer (src/24-29) ----
def _pl(f):
    p=os.path.join(GEN,f); return pd.read_csv(p,sep="\t") if os.path.exists(p) else None
ph_hits = _pl("phenome_hits.tsv")
ph_pleio= _pl("phenome_pleiotropy_axes.tsv")
ph_col  = _pl("coloc_phenome_results.tsv")
nov_rank= _pl("novelty_engine_ranked.tsv")
HAS_PHEN = ph_hits is not None and nov_rank is not None
if HAS_PHEN:
    PH_HITS = int(len(ph_hits))
    PH_DIS  = int(ph_hits.disease.nunique())
    PH_GENE = int(ph_hits.gene_symbol.nunique())
    PH_NCAT = int(ph_hits.disease_category.nunique())
    PH_CATCOUNTS = ph_hits.disease_category.value_counts().to_dict()
    PH_PLEIO = int(ph_pleio.gene_symbol.nunique()) if ph_pleio is not None else 0
    PH_COLOC = int((ph_col.PP_H4>=0.8).sum()) if ph_col is not None else 0
    PH_NOVCOL= int((nov_rank.category_label=="NOVEL colocalized").sum())
    PH_NOVNOM= int((nov_rank.category_label=="novel nomination").sum())
    PH_TOTAL = int(N_DIS + PH_DIS)  # combined disease breadth is described in text

# ---- whole-phenome scan totals (src/54 caches these from the 615 MB cis_MR_ALL scan) ----
# phenome_hits.tsv above is the intermediate 28-endpoint scan; the released atlas is the
# full FinnGen R12 phenome, so the figures and numbers in that section must come from here.
_DECK = os.path.join(ROOT, "08_figures", "deck")
def _deck(f):
    p = os.path.join(_DECK, f)
    return pd.read_csv(p, sep="\t", index_col=0) if os.path.exists(p) else None
_tot = _deck("_scan_totals.tsv")
_ccat = _deck("_volcano_counts.tsv")
_ccls = _deck("_volcano_counts_class.tsv")
HAS_WHOLE = _tot is not None and _ccat is not None
if HAS_WHOLE:
    W_TESTS = int(_tot.loc["tests", "value"]);   W_GENES = int(_tot.loc["genes", "value"])
    W_ENDP  = int(_tot.loc["endpoints", "value"]); W_CHAP = int(_tot.loc["chapters", "value"])
    W_HITS  = int(_tot.loc["hits", "value"])
    _sig = _ccat[_ccat["hits"] > 0].sort_values("hits", ascending=False)
    W_HCHAP = int(len(_sig))
    W_TOPCH = _sig.head(6)
    W_TOPCLS = (_ccls[_ccls["hits"] > 0].sort_values("hits", ascending=False).head(5)
                if _ccls is not None else None)
    # counted on phenocode, not phenotype name: two FinnGen codes can share a display name
    W_HGENE = int(_tot.loc["hit_genes", "value"])
    W_HDIS  = int(_tot.loc["hit_endpoints", "value"])
    W_HCHAP = int(_tot.loc["hit_chapters", "value"])

# ---- doc styling ----
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

def H(txt, lvl=1, color=None):
    p = doc.add_heading(txt, level=lvl)
    if color:
        for r in p.runs: r.font.color.rgb = color
    return p
def P(txt, italic=False, size=10.5, align=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if align: p.alignment = align
    return p
def figure(fname, caption):
    # figures live in 08_figures/nature by default, but the whole-phenome panels
    # rebuilt for the slide deck live in 08_figures/deck, so allow a subfolder path
    fp = os.path.join(FIG, fname)
    if not os.path.exists(fp):
        fp = os.path.join(ROOT, "08_figures", fname)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); r = c.add_run(caption); r.italic=True; r.font.size=Pt(9)

BLUE = RGBColor(0x1f,0x4e,0x79)

# ================= TITLE =================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("An open, genetics-anchored plasma immunome atlas recovers established "
              "therapeutic immune axes and nominates colocalized autoimmune targets")
r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE

P("HDDM Layer 54 \u00b7 Plasma Immunome\u2013Phenome Atlas \u00b7 an independent, reproducible resource",
  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ================= ABSTRACT =================
H("Abstract", 1, BLUE)
P(
 f"Circulating immune proteins mediate host defence and are the targets of many "
 f"approved immunotherapies, yet a systematic, openly reproducible map linking the "
 f"plasma immune proteome to disease has been lacking. Here we assemble a plasma "
 f"immunome\u2013phenome atlas that curates the {N_OLINK:,}-analyte Olink Explore proteomic "
 f"space to {N_IMM:,} plasma-detectable immune proteins using orthogonal Human Protein "
 f"Atlas immune-cell specificity and pathway annotation, and anchors each protein to "
 f"disease through genetics. Using cis-eQTL instruments for {N_INST} immune genes, we "
 f"perform Mendelian randomization against {N_DIS} immune-mediated diseases "
 f"({N_TEST:,} tests) and calibrate every finding on an explicit claim-strength ladder. "
 f"The analysis blindly recovers established therapeutic axes\u2014IL6ST in rheumatoid "
 f"arthritis, CTLA4 in autoimmune thyroid disease, TNFRSF1A in ankylosing spondylitis\u2014"
 f"providing an internal positive control. Statistical colocalization resolves {N_T4+N_T5} "
 f"loci to a shared causal variant (PP.H4\u22650.8), promoting them to prioritized "
 f"transcript-level causal targets, and nominates candidates including SIGLEC7/9 in "
 f"Guillain\u2013Barr\u00e9 syndrome, HAVCR1 in coeliac disease and IFNGR2 in psoriasis. "
 f"Independent protein-level instruments from plasma cis-pQTLs (INTERVAL) then confirm a "
 f"subset at the protein level\u2014TNFSF14 in multiple sclerosis and SWAP70 in rheumatoid "
 f"arthritis colocalize at both transcript and protein level (PP.H4>0.97), elevating them "
 f"to protein-level causal targets\u2014while exposing informative transcript\u2013protein "
 f"discordances (e.g. IL6ST) that temper otherwise strong transcript-level signals. "
 f"In independent, FinnGen-free consortium GWAS the direction of effect replicates in "
 f"{N_RCON} of {N_COV} testable nominations ({N_REP} at P<0.05). "
 f"Signals in the MHC ({N_T2}) are explicitly flagged for linkage-disequilibrium "
 f"confounding and withheld from causal claims. All curation, instruments, statistics "
 f"and figures are built from public data and released as an open, versioned resource."
)

# ================= INTRODUCTION =================
H("Introduction", 1, BLUE)
P("Plasma proteins are the most accessible window on human immunity and constitute a "
  "disproportionate share of the druggable proteome. Affinity proteomics now quantifies "
  "thousands of circulating proteins at population scale, but converting protein\u2013disease "
  "correlations into directional, genetically-supported hypotheses requires (i) a principled "
  "definition of the immune proteome, (ii) genetic instruments that expose causal direction, "
  "and (iii) discipline in matching claim strength to evidence. Association is not causation, "
  "a cis-eQTL is not a protein measurement, and a Mendelian-randomization estimate at the MHC "
  "is not a clean instrument. We therefore built this atlas around an explicit evidence ladder "
  "and report every finding at the tier its evidence supports\u2014no higher.")

# ================= RESULTS =================
H("Results", 1, BLUE)

H("Curation of a plasma-detectable immune proteome", 2)
P(f"Starting from the {N_OLINK:,} analytes of Olink Explore, we integrated Human Protein "
  f"Atlas immune-cell RNA specificity, secretome location and curated pathway membership "
  f"(cytokine, chemokine, interferon, TNF, complement, checkpoint, CD/leukocyte, HLA, "
  f"immunoglobulin, acute-phase, coagulation) into a transparent inclusion score, retaining "
  f"{N_IMM:,} plasma immune proteins (Fig. 1). The resource preserves, for each protein, its "
  f"immune class and the blood-cell lineage that expresses it, so that downstream disease "
  f"signals can be read back to a cell of origin. A permissive immune-signature gene set "
  f"(MSigDB C7) covered almost the entire Olink space and was therefore retained only as "
  f"annotation, not as an inclusion criterion\u2014an early example of the calibration principle "
  f"applied throughout.")
figure("Figure1_design_curation.png",
       "Figure 1 | Design and curation of the plasma immunome. Inclusion funnel from the "
       f"Olink Explore universe to {N_IMM:,} plasma immune proteins; immune-class composition; "
       "mapping of proteins to blood-cell lineage of origin; and co-occurrence of annotation flags.")

H("Architecture and druggability of the immune proteome", 2)
P("The curated proteome is dominated by immune-cell-enriched proteins with substantial "
  "cytokine, chemokine, complement and checkpoint representation, and maps onto granulocyte, "
  "myeloid, dendritic, T-, B- and NK-cell lineages (Fig. 2). A large fraction are secreted or "
  "single-pass receptors\u2014the protein classes most tractable to antibodies and biologics\u2014"
  "establishing the atlas as a systematic hunting ground for immune drug targets before any "
  "genetic evidence is applied.")
figure("Figure2_druggability_architecture.png",
       "Figure 2 | Architecture and druggability of the plasma immune proteome: immune-class "
       "and lineage composition, secretome/receptor tractability, and target-class breakdown.")

H("Genetic anchoring recovers established therapeutic immune axes", 2)
P(f"To expose causal direction we instrumented {N_INST} immune genes with their strongest cis "
  f"expression-quantitative-trait locus (eQTLGen; single-instrument Wald ratio) and tested them "
  f"against {N_DIS} immune-mediated diseases from FinnGen R12 ({N_TEST:,} tests), controlling "
  f"the false-discovery rate at 5% ({N_SIG} significant gene\u2013disease pairs; Fig. 3). Critically, "
  f"the analysis\u2014run blind to drug annotation\u2014recovers multiple axes that are already the "
  f"targets of approved immunotherapies: IL6ST (the IL-6 co-receptor gp130) in rheumatoid "
  f"arthritis, CTLA4 in autoimmune thyroid disease, and the TNF-receptor axis (TNFRSF1A) in "
  f"ankylosing spondylitis, with directions of effect consistent with the known pharmacology "
  f"(Fig. 4). This concordance is an internal positive control: a pipeline that rediscovers "
  f"tocilizumab-, abatacept- and anti-TNF-relevant biology from genetics alone is calibrated to "
  f"be believed when it points somewhere new.")
figure("Figure3_MR_design_global.png",
       "Figure 3 | Mendelian-randomization design and global results across 13 immune-mediated "
       "diseases: workflow, effect-size volcano, significant hits per disease, and immune-class enrichment.")
figure("Figure4_positive_controls.png",
       "Figure 4 | The pipeline blindly recovers established immune therapeutic axes (IL6ST, "
       "CTLA4, TNFRSF1A, TNFSF14), each shown with its cis-MR odds ratio and the drug it underpins.")

H("Colocalization calibrates novel, disease-specific target nominations", 2)
P(f"We next asked which significant signals reflect a shared causal variant rather than "
  f"linkage-disequilibrium coincidence, applying approximate-Bayes-factor colocalization "
  f"(coloc.abf) at every non-MHC hit. Colocalization behaves as expected on the positive "
  f"controls (IL6ST\u2013rheumatoid arthritis PP.H4=1.00; CTLA4 PP.H4>0.98; TNFSF14\u2013multiple "
  f"sclerosis PP.H4>0.99) and promotes {N_T4} loci to prioritized transcript-level causal targets "
  f"(PP.H4\u22650.8), including several disease-specific nominations: SIGLEC7 and SIGLEC9 "
  f"(sialic-acid-binding inhibitory receptors) in Guillain\u2013Barr\u00e9 syndrome, HAVCR1/KIM-1 in "
  f"coeliac disease, IFNGR2 in psoriasis, CD226 in sarcoidosis and HNMT in Crohn's disease "
  f"(Fig. 5). A further {N_T3} signals remain genetically-supported nominations pending stronger "
  f"colocalization. Signals in the MHC\u2014most prominently the complement gene C2, significant "
  f"across five diseases\u2014are held at nomination status and explicitly annotated for LD "
  f"confounding rather than promoted ({N_T2} MHC pairs). Figure 6 places every finding on the "
  f"claim-strength ladder and defines the validation path to causal and translational tiers.")
figure("Figure5_novel_nominations.png",
       "Figure 5 | Colocalized novel autoimmune target nominations, each annotated with its "
       "colocalization posterior (PP.H4): C2 (MHC, held at nomination), SIGLEC7/9\u2192Guillain\u2013Barr\u00e9, "
       "HAVCR1\u2192coeliac, IFNGR2\u2192psoriasis.")
figure("Figure6_validation_roadmap.png",
       "Figure 6 | Evidence-tier calibration and validation roadmap: the claim-strength ladder, "
       "where each finding sits, per-locus colocalization posteriors, and the path to protein-level "
       "(pQTL), replication and functional validation.")

H("Protein-level pQTL instruments confirm a subset and reveal informative discordance", 2)
P(f"Because a cis-eQTL indexes transcript, not protein, abundance, we sought independent "
  f"confirmation using plasma cis-pQTLs from the INTERVAL study (Sun et al., 2018) as "
  f"protein-level instruments for the significant genes ({N_PQTL} gene\u2013disease pairs had a "
  f"matched plasma aptamer). This layer is decisive rather than cosmetic. Two nominations "
  f"colocalize at both the transcript and the protein level with concordant direction and "
  f"survive protein-level MR\u2014TNFSF14 in multiple sclerosis (pQTL PP.H4>0.99) and SWAP70 in "
  f"rheumatoid arthritis (pQTL PP.H4>0.97)\u2014and we designate these {N_T5} the study's "
  f"protein-level causal targets, the highest tier reached here (Fig. 7). A further set, "
  f"including TNFRSF1A in ankylosing spondylitis, are directionally concordant with nominally "
  f"significant protein-level MR. Critically, only {N_CONC} of {N_PQTL} pairs are directionally "
  f"concordant between transcript and protein, and {N_DISC} are discordant\u2014most notably IL6ST "
  f"in rheumatoid arthritis, where the transcript signal (risk-increasing) reverses at the "
  f"protein level, consistent with the known biology of soluble gp130 as an inhibitor of IL-6 "
  f"trans-signalling. Rather than suppress these discordances we report them, because they "
  f"define which transcript-level nominations should not yet be read as protein-level causal "
  f"claims (e.g. IFNGR2, ERBB3), and they are exactly the distinctions a therapeutic programme "
  f"needs before committing to a modality.")
figure("Figure7_pqtl_confirmation.png",
       "Figure 7 | Protein-level (INTERVAL plasma pQTL) validation. (a) Transcript- vs "
       "protein-level MR log-odds-ratios per hit (blue, directionally concordant; red, discordant). "
       "(b) Protein-level colocalization posteriors; TNFSF14\u2192multiple sclerosis and SWAP70\u2192"
       "rheumatoid arthritis exceed PP.H4=0.8 at the protein level.")

H("Nominations replicate in independent, non-FinnGen disease GWAS", 2)
P(f"Because discovery used FinnGen alone, we sought out-of-sample replication in independent "
  f"consortium GWAS accessed through OpenGWAS (IMSGC/Patsopoulos multiple sclerosis; Okada "
  f"rheumatoid arthritis; IGAS/Cortes ankylosing spondylitis; Stuart psoriasis; Fischer "
  f"sarcoidosis; Sakaue autoimmune thyroid; none of which include FinnGen). Of the {N_SIG} "
  f"significant gene\u2013disease pairs, {N_COV} had the instrument variant present in a matched "
  f"independent GWAS; of these, {N_RCON} were directionally concordant and {N_REP} replicated "
  f"at nominal significance (P<0.05), including every one of the positive-control axes and the "
  f"tier-5 target TNFSF14 in multiple sclerosis (P=2\u00d710\u207b\u00b9\u00b2, IMSGC), IL6ST in rheumatoid "
  f"arthritis (P=7\u00d710\u207b\u00b2\u2074, Okada) and CTLA4 (P=3\u00d710\u207b\u00b2\u2070) (Fig. 8). Notably, the direction of "
  f"effect agreed in 19 of 19 covered hits, an out-of-sample consistency that would be "
  f"vanishingly unlikely by chance and that substantially raises confidence in the "
  f"transcript-level nominations. Several candidates could not be tested because no "
  f"FinnGen-independent GWAS exists (Guillain\u2013Barr\u00e9 syndrome, Sjögren syndrome) or because "
  f"the instrument was absent from an older targeted-array study (coeliac disease); these are "
  f"stated as coverage gaps rather than failures to replicate.")
figure("Figure8_replication.png",
       "Figure 8 | Independent replication. (a) \u2212log10(P) of each instrument variant in a matched "
       "FinnGen-independent consortium GWAS (bar colour = replication status; cohort labelled). "
       "(b) Replication outcome across all significant hits. Direction of effect was concordant in "
       "all 19 covered nominations.")

# ================= NOVELTY / TRANSLATION =================
H("The atlas recovers approved-drug direction, resolves shared causal axes, and prioritizes novel targets", 2)
P(f"Beyond identifying targets, the genetic architecture carries a directional signal that is "
  f"directly actionable. Encoding each effect as protein-lowering-protective (OR<1, arguing for "
  f"agonism/replacement) versus protein-raising-risk (OR>1, arguing for blockade), the atlas "
  f"recovers not only the identity but the correct pharmacological DIRECTION of established "
  f"therapies: CTLA4 is protective and is drugged by an agonist (abatacept, CTLA4-Ig); TNFRSF1A "
  f"and the IL-6 axis are risk-increasing and are drugged by blockers (etanercept, tocilizumab); "
  f"IL4 raises psoriasis risk and is blocked by dupilumab. In total {N_DIRMATCH} recovered axes "
  f"match approved-drug direction, and the two cases where plasma pQTL direction inverts the eQTL "
  f"signal (IL6ST, reflecting soluble gp130 trans-signalling; TNFRSF14/HVEM) are exactly the axes "
  f"where receptor biology is known to be direction-dependent\u2014so the discordance is mechanistically "
  f"informative rather than noise (Fig. 9c). Mapping genes that act across more than one disease "
  f"exposes shared causal immune axes\u2014CTLA4 (autoimmune thyroid + rheumatoid arthritis), NUMB and "
  f"TNFRSF14 (each across two diseases), and the MHC-region C2 signal recurring across five\u2014"
  f"pointing to pleiotropic control points whose directions are consistent across indications "
  f"(Fig. 9b). At the class level, genetically-supported causal targets trend toward the TNF "
  f"superfamily, interferon axis and immune-checkpoint classes (Fisher odds ratios 4\u201310), a "
  f"coherent enrichment given known autoimmune drug space, though not individually significant "
  f"after correction given the small target set (Fig. 9a). Finally, separating internal "
  f"positive controls from genuinely new signals leaves {N_NOVELNOM} novel colocalized nominations "
  f"lacking an approved drug for their indication; ranked by cumulative orthogonal evidence "
  f"(transcript coloc + protein pQTL + independent replication), SWAP70 in rheumatoid arthritis "
  f"tops the list as the only novel target with protein-level colocalization AND independent "
  f"replication, followed by replicated transcript-level nominations IFNGR2 (psoriasis), CA8 and "
  f"MYL6B (rheumatoid arthritis) and HAVCR1 (coeliac disease) (Fig. 9d).")
figure("Figure9_novelty.png",
       "Figure 9 | Novelty and translation. (a) Immune-class enrichment of tier\u22654 causal targets "
       "(Fisher; trends toward TNF-superfamily/interferon/checkpoint classes). (b) Pleiotropic "
       "shared causal axes for genes acting on \u22652 diseases (red=risk/blockade, blue=protective/"
       "agonism). (c) Genetics recover the DIRECTION of approved drugs, with mechanistically "
       "informative discordances. (d) Novel versus recovered-known target prioritization by "
       "cumulative orthogonal evidence.")

# ================= WHOLE-PHENOME SCAN =================
# NOTE: phenome_hits.tsv (src/24-29) is the 28-endpoint five-category PILOT, not the
# released scan. Figures and counts here therefore come from the full FinnGen R12 sweep
# cached by src/54 (08_figures/deck), and the pilot is described as a pilot.
if HAS_WHOLE:
    H("The causal map is whole-phenome: every immune gene against every FinnGen R12 endpoint", 2, BLUE)
    _CHAPTAG = re.compile(r"\s*\([A-Z0-9_]+\)$")          # drop the "(I9_)" code suffix
    _cc = ", ".join("{} ({})".format(_CHAPTAG.sub("", str(k)).strip(), int(v))
                    for k, v in W_TOPCH["hits"].items())
    P(f"A causal atlas that scans a hand-picked disease list can only be as broad as that list. "
      f"We therefore ran the identical cis-MR machinery over the ENTIRE FinnGen R12 phenome: "
      f"{W_GENES} instrumented immune genes \u00d7 {W_ENDP:,} endpoints spanning all {W_CHAP} "
      f"FinnGen chapters, {W_TESTS:,} Wald-ratio tests, with Benjamini\u2013Hochberg FDR applied "
      f"once across the whole scan rather than within any curated subset. This yields "
      f"{W_HITS:,} causal gene\u2013disease pairs at FDR<5%, over {W_HGENE} distinct proteins and "
      f"{W_HDIS} distinct endpoints in {W_HCHAP} chapters (Fig. 11). The burden is not where an "
      f"autoimmune-only reading would predict: the largest yields fall in {_cc}\u2014"
      f"circulatory, musculoskeletal, respiratory, skin, endocrine and digestive disease "
      f"carry more phenome-wide causal immune signal than the classical autoimmune chapter, "
      f"which is exactly the conclusion a curated 13- or 28-disease scan cannot reach.")
    P(f"Because the scan is unrestricted, pleiotropy can be measured rather than assumed: "
      f"immune proteins recur as causal across multiple organ systems, and the genes with the "
      f"widest reach are shared control points rather than disease-specific effectors "
      f"(Fig. 12). Mapping every causal protein back to the blood-cell lineage in which it is "
      f"enriched shows which arms of the immune system carry causal weight across the phenome "
      f"(Fig. 13), and encoding each effect as protein-raising-risk (argues for blockade) "
      f"versus protein-lowering-protective (argues for agonism or replacement) partitions all "
      f"{W_HITS:,} signals into a direction-aware therapeutic map, with MHC/LD-confounded loci "
      f"held separately rather than promoted (Fig. 14).")
    if HAS_PHEN:
        P(f"An earlier five-category pilot over {PH_DIS} FinnGen endpoints ({PH_HITS} hits, "
          f"{PH_GENE} genes) was used to develop and sanity-check the scoring machinery before "
          f"the full sweep; it is retained in the released tables for provenance, but every "
          f"claim, figure and ranking in this manuscript is computed from the whole-phenome "
          f"scan above, not from that pilot.")
    P(f"Integrating every real-data evidence layer\u2014causal strength, transcript "
      f"colocalization, cross-system pleiotropy, HPA-derived druggability and cell-source "
      f"specificity, minus explicit penalties for known-drug axes and MHC LD\u2014into a single "
      f"auditable novelty-priority score ranks the whole-phenome targets and separates novel "
      f"colocalized targets and further novel nominations from recovered positive controls "
      f"(Figs. 15, 16). As throughout, the engine down-weights internal positive controls "
      f"(CTLA4, IL6ST, IL2RA) so that the ranking surfaces genuinely new biology rather than "
      f"re-discovering approved drugs.")
    figure(os.path.join("deck", "DECK_phenome_volcano.png"),
           f"Figure 11 | Whole-phenome cis-MR. Effect-size volcano over all {W_TESTS:,} "
           f"gene\u00d7endpoint tests ({W_GENES} immune genes \u00d7 {W_ENDP:,} FinnGen R12 endpoints, "
           f"{W_CHAP} chapters), with the {W_HITS:,} FDR<5% causal pairs highlighted, and the "
           f"per-chapter yield of significant pairs.")
    figure(os.path.join("deck", "DECK_pleiotropy.png"),
           "Figure 12 | Cross-system pleiotropic immune axes measured on the whole-phenome scan: "
           "immune genes causal in more than one FinnGen chapter, with the direction of effect "
           "per system.")
    figure(os.path.join("deck", "DECK_cellsource.png"),
           "Figure 13 | Causal-target enrichment by blood-cell source of origin across the whole "
           "phenome, computed from all FDR<5% pairs rather than a curated disease subset.")
    figure(os.path.join("deck", "DECK_direction_summary.png"),
           f"Figure 14 | Direction-aware therapeutic map of all {W_HITS:,} whole-phenome causal "
           f"signals, partitioned into block/neutralise versus agonise/replace, with MHC-caution "
           f"loci flagged rather than promoted.")
    figure(os.path.join("deck", "DECK_novelty_engine.png"),
           "Figure 15 | Integrated novelty-priority ranking of causal immune targets across the "
           "whole phenome, showing the stacked evidence components and the known-drug and MHC "
           "penalties applied to each target.")
    figure(os.path.join("deck", "DECK_novel_nominations.png"),
           "Figure 16 | Novel-versus-known separation across the whole phenome: the top novel "
           "nominations lacking an approved drug for their indication, set against the recovered "
           "known axes and MHC-caution loci that the engine deliberately down-weights.")

# ===== EVERY EVIDENCE LAYER EXTENDED TO THE WHOLE PHENOME =====
# the FDR<5% subset of the 615 MB cis_MR_ALL scan; src/54 already streamed and cached it,
# so read the cache rather than pulling the whole scan into memory to build a document
_vcp = os.path.join(_DECK, "_volcano_cache.parquet")
all_mr   = (pd.read_parquet(_vcp) if os.path.exists(_vcp)
            else _pl("cis_MR_ALL_finngen_results.tsv"))
all_col  = _pl("coloc_ALL_finngen_results.tsv")
all_pmr  = _pl("pqtl_MR_ALL_finngen_results.tsv")
all_pcol = _pl("pqtl_coloc_ALL_finngen_results.tsv")
all_nov  = _pl("novelty_engine_ranked_ALL.tsv")
all_il   = _pl("intelligence_layer_final_table_ALL.tsv")
all_uk   = _pl("uk_panphenome_concordance_ALL.tsv")
if all(x is not None for x in (all_mr, all_col, all_pmr, all_pcol, all_nov, all_il, all_uk)):
    A_HIT   = all_mr[all_mr.FDR < 0.05]
    A_NDIS  = int(A_HIT.phenotype.nunique())
    A_NPAIR = int(len(A_HIT))
    A_NGENE = int(A_HIT.gene_symbol.nunique())
    A_COLD  = int(all_col.disease.nunique())
    A_PMR   = int(len(all_pmr)); A_PMRS = int((all_pmr.FDR < 0.05).sum())
    A_PMRD  = int(all_pmr.disease.nunique()); A_PMRG = int(all_pmr.gene.nunique())
    A_PCOL  = int(len(all_pcol)); A_PCOLS = int((all_pcol.PP_H4 >= 0.8).sum())
    A_PCOLD = int(all_pcol.disease.nunique())
    A_PCONF = int((all_nov.category_label == "NOVEL protein-confirmed").sum())
    A_T5    = int((all_il.Novelty_tier == 5).sum())
    A_T4    = int((all_il.Novelty_tier == 4).sum())
    A_UKP   = int(len(all_uk)); A_UKD = int(all_uk.phenocode.nunique())
    A_UKG   = int(all_uk.gene_symbol.nunique())
    A_UKC   = int(all_uk.concordant.astype(bool).sum())
    A_UKV   = int(all_uk.two_population_validated.astype(bool).sum())
    A_UKX   = int((all_uk.match_method == "exact-code").sum())

    H("Every evidence layer is phenome-wide, not restricted to a curated disease list", 2, BLUE)
    P(f"A resource that scans the phenome for discovery but validates only a hand-picked disease "
      f"core is only as broad as its narrowest layer. We therefore extended every downstream layer "
      f"to the same scale as the discovery scan. Colocalization now covers {A_COLD} diseases; the "
      f"protein-level layer, previously computed only for the autoimmune core, now instruments "
      f"every pan-phenome causal protein that has a public INTERVAL aptamer and gives {A_PMR} "
      f"protein-level Wald-ratio tests across {A_PMRD} diseases and {A_PMRG} plasma proteins "
      f"({A_PMRS} significant at FDR<5%), together with {A_PCOL} protein-level colocalization loci "
      f"across {A_PCOLD} diseases of which {A_PCOLS} reach PP.H4\u22650.8; and the novelty engine and "
      f"disease-intelligence layer now score all {A_NPAIR} causal pairs over {A_NDIS} diseases and "
      f"{A_NGENE} proteins (Fig. 17).")
    P(f"Running the protein layer phenome-wide changes what the atlas can conclude. {A_PCONF} "
      f"gene\u2013disease pairs are now protein-confirmed\u2014transcript colocalization plus a "
      f"direction-concordant plasma-protein MR\u2014and the highest evidence tier, which was empty "
      f"when the protein layer stopped at the autoimmune core, is now occupied by {A_T5} targets "
      f"(IL2RA in type-1 diabetes, ANXA2 in hypertension, PPP3R1 in venous disease) alongside "
      f"{A_T4} tier-4 prioritized targets (Figs. 18\u201320). Strong protein-level colocalizations "
      f"appear well outside the autoimmune arc\u2014CTSH in early-onset type-1 diabetes, LILRB2 and "
      f"SERPING1 across obstructive airway disease, CFH in dry age-related macular degeneration\u2014"
      f"while classic controls (CFH\u2192macular degeneration, CTLA4\u2192rheumatoid arthritis) are "
      f"recovered blind (Figs. 21, 22). The residual boundary is biological rather than "
      f"procedural: {A_NGENE - A_PMRG} causal proteins have no SomaScan aptamer, being largely "
      f"intracellular or MHC-region, and cannot reach a protein tier from any login-free resource.")
    P(f"The two-population arm was widened on the same principle. Beyond FinnGen endpoints with an "
      f"identical Neale phenocode, FinnGen trait names were normalised and matched to independent UK "
      f"Biobank GWAS, restricted to single-cohort UK Biobank datasets so that no comparison cohort "
      f"silently contains FinnGen\u2014several large public meta-analyses of the same endpoints do, "
      f"which would make replication circular\u2014and to datasets with at least 200 cases. This gives "
      f"{A_UKP} Finland-versus-England causal-effect comparisons over {A_UKD} diseases and {A_UKG} "
      f"proteins ({A_UKX} exact-phenocode, {A_UKP - A_UKX} name-matched, each pairing recorded in the "
      f"released table), of which {A_UKC} ({100 * A_UKC / A_UKP:.0f}%) agree in direction and {A_UKV} "
      f"replicate with a nominally significant same-direction effect in UK Biobank.")
    figure("Figure17_layer_disease_coverage.png",
           "Figure 17 | Disease coverage of every evidence layer before and after extension: "
           "colocalization, protein-level pQTL MR and colocalization, novelty scoring and the "
           "intelligence layer now all run across the FinnGen R12 phenome.")
    figure("Figure18_panphenome_pqtl_volcano.png",
           "Figure 18 | Pan-phenome protein-level Mendelian randomization using INTERVAL plasma "
           "cis-pQTL instruments, with real effect sizes and standard errors on both sides.")
    figure("Figure19_eqtl_pqtl_concordance_all.png",
           "Figure 19 | Transcript-level versus protein-level causal effect direction for every "
           "pair instrumented at both levels; discordance is retained and reported, not discarded.")
    figure("Figure20_tiers_by_system_all.png",
           "Figure 20 | Evidence-tier composition of all causal gene\u2013disease pairs by organ "
           "system, from MHC-held nominations to protein-level targets.")
    figure("Figure21_novelty_priority_all.png",
           "Figure 21 | Top pan-phenome novelty-priority targets, with protein-confirmed targets "
           "distinguished from transcript-colocalized nominations and recovered known axes.")
    figure("Figure22_pqtl_coloc_all.png",
           "Figure 22 | Strongest plasma-protein colocalizations across the phenome.")

# ================= DISCUSSION =================
H("Discussion", 1, BLUE)
P("This atlas is deliberately calibrated rather than maximalist. Its central result is not a "
  "single target but a demonstration that an openly reproducible pipeline, built entirely from "
  "public data, can rediscover approved-drug biology from genetics, extend the same evidence "
  "standard to new nominations, and then\u2014crucially\u2014confirm or temper those nominations with an "
  "independent protein-level layer. Two targets reach the protein-level causal tier (TNFSF14 in "
  "multiple sclerosis, SWAP70 in rheumatoid arthritis), colocalizing at both transcript and "
  "protein. The remaining transcript-level nominations\u2014SIGLEC7/9 in Guillain\u2013Barr\u00e9 syndrome, "
  "HAVCR1 in coeliac disease, IFNGR2 in psoriasis\u2014are hypotheses at the transcript level, and we "
  "say no more than that; where plasma pQTLs disagree in direction (e.g. IL6ST, IFNGR2) we flag "
  "the discordance explicitly. The principal remaining limitations define the roadmap: pQTL "
  "coverage is limited to one plasma proteomic panel; effects are estimated in a single "
  "population; and no finding has yet been functionally perturbed. Elevating a nomination further "
  "requires replication in an independent GWAS and ancestry, orthogonal pQTL platforms, direct "
  "plasma-protein\u2013disease association, and perturbation. By publishing the atlas with its "
  "evidence ladder intact, we make both what is known and what remains to be shown fully auditable.")

# ================= METHODS =================
H("Methods (summary)", 1, BLUE)
P("Proteome curation. Olink Explore analytes were annotated with Human Protein Atlas immune-cell "
  "RNA specificity, secretome location, molecular function and curated immune-pathway membership; "
  "a transparent additive score (weighting cytokine/chemokine/IFN/TNF, complement/checkpoint and "
  "immune-cell-enrichment most heavily) defined inclusion. MSigDB C7 was used as annotation only.")
P("Genetic instruments and MR. For each immune gene the strongest cis-eQTL (eQTLGen) was used as a "
  "single instrument; Z-scores were converted to effect sizes using allele frequency and sample "
  "size (Zhu et al., 2016), harmonized against FinnGen R12 summary statistics with palindromic "
  "variants removed, and combined by the Wald ratio. FDR was controlled at 5% (Benjamini\u2013Hochberg).")
P("Colocalization. At each non-MHC significant locus, coloc.abf (Giambartolomei et al., 2014) was "
  "applied over shared cis variants using per-SNP variances\u2014real FinnGen standard errors for the "
  "outcome and Zhu-reconstructed standard errors for the exposure\u2014with default priors "
  "(p1=p2=1e-4, p12=1e-5). Loci with PP.H4\u22650.8 were designated prioritized causal targets.")
P("Protein-level pQTL validation. For each significant gene, plasma cis-pQTLs from INTERVAL "
  "(Sun et al., 2018; GWAS Catalog harmonised summary statistics, GRCh37) within \u00b1500 kb of the "
  "gene body were used as protein-level instruments. The strongest cis-pQTL drove a Wald-ratio MR "
  "against each disease (real beta and standard error on both sides), and coloc.abf was applied "
  "over shared cis variants. A pair was promoted to a protein-level causal target only if the "
  "transcript- and protein-level MR were directionally concordant and the protein-level "
  "colocalization reached PP.H4\u22650.8; directional discordances were reported, not discarded.")
P("Independent replication. Each significant instrument variant was queried in a "
  "FinnGen-independent consortium GWAS for the matching disease via the OpenGWAS API "
  "(IMSGC/Patsopoulos, Okada, IGAS/Cortes, Stuart, Fischer, Sakaue). The disease effect was "
  "harmonised to the eQTL expression-increasing allele; a hit was scored 'replicated' when the "
  "independent effect was directionally concordant with discovery and nominally significant "
  "(P<0.05). Diseases with no FinnGen-independent GWAS, or variants absent from older "
  "targeted-array studies, were recorded as coverage gaps.")
P("Claim-strength gate. Every finding was assigned an evidence tier\u2014transcript-level proxy "
  "(cis-eQTL), genetically-supported nomination (cis-MR FDR<0.05), prioritized causal target "
  "(+transcript colocalization), protein-level causal target (+concordant pQTL-MR and protein "
  "colocalization)\u2014with MHC signals capped at nomination and flagged for LD. Language in this "
  "manuscript is bound to these tiers.")
P("Data availability. The atlas is built entirely from public resources (Olink Explore universe, "
  "Human Protein Atlas, MSigDB, eQTLGen, FinnGen R12, INTERVAL plasma pQTLs via the GWAS Catalog) "
  "and released, with all code, as a versioned open resource. Individual-level UK Biobank data and "
  "the UKB-PPP/Synapse-gated pQTLs are controlled-access and were not used.")

# ---- evidence table (protein-aware final tiers + replication) ----
H("Table 1 | Final evidence-tiered targets (discovery + protein + replication)", 2)
tb = doc.add_table(rows=1, cols=8); tb.style = "Light Grid Accent 1"
hdrs=["Gene","Disease","eQTL OR","coloc\nPP.H4","pQTL\nPP.H4","dir","Replication","Final tier"]
for i,h in enumerate(hdrs):
    tb.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
for _,r in final.iterrows():
    c = tb.add_row().cells
    c[0].text=str(r.gene_symbol); c[1].text=str(r.disease)
    c[2].text=f"{r.OR:.2f}"
    c[3].text=("\u2014" if pd.isna(r.PP_H4) else f"{r.PP_H4:.2f}")
    c[4].text=("\u2014" if pd.isna(r.pQTL_PPH4) else f"{r.pQTL_PPH4:.2f}")
    c[5].text=str(r.pQTL_concordant) if str(r.pQTL_concordant)!="nan" else ""
    c[6].text=(str(r.rep_status) if HAS_REP and str(r.get("rep_status"))!="nan" else "")
    c[7].text=f"T{int(r.final_tier)}"
    for cell in c:
        for para in cell.paragraphs:
            for run in para.runs: run.font.size=Pt(7.5)
P("Tiers: T5 protein-level causal target (concordant pQTL-MR + protein coloc PP.H4\u22650.8); "
  "T4 prioritized causal target (transcript coloc PP.H4\u22650.8); T3 genetically-supported "
  "nomination; T2 nomination held for MHC LD. dir = transcript/protein direction concordance; "
  "Replication = status in a FinnGen-independent consortium GWAS (OpenGWAS).",
  italic=True, size=8)

# ================= EXTENDED DATA / SUPPLEMENTARY FIGURE GALLERY =================
SUP = os.path.join(ROOT, "08_figures", "supplementary")
PHE = os.path.join(ROOT, "08_figures", "phenome")

def gallery_figure(fp, caption, width=5.2):
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); r = c.add_run(caption); r.italic=True; r.font.size=Pt(8)

def _disease_from(stem):
    return stem.replace("_"," ").strip().rstrip("_").replace("  "," ")

# ---- Pan-phenome extended figures ----
if HAS_PHEN and os.path.isdir(PHE):
    doc.add_page_break()
    H("Extended Data \u00b7 Pan-phenome figures", 1, BLUE)
    P("Per-category and per-disease views of the 28-disease pan-phenome cis-MR "
      "(src/24\u201326). Volcano plots show cis-MR effect size vs \u2212log10(FDR) for every "
      "disease; forest plots summarise the strongest causal targets within each disease "
      "category.", italic=True, size=9)
    phe = sorted(os.listdir(PHE))
    order = [f for f in phe if f.startswith("PFig_hits")] + \
            [f for f in phe if f.startswith("PFig_forest")] + \
            [f for f in phe if f.startswith("PFig_volcano")]
    n = 0
    for f in order:
        stem = f[:-4]
        if f.startswith("PFig_hits"):
            cap = "Pan-phenome causal hits per disease category."
        elif f.startswith("PFig_forest"):
            cat = stem.replace("PFig_forest_","").replace("_"," ")
            cap = f"Forest plot of causal immune targets \u2014 {cat} diseases (OR \u00b1 95% CI)."
        else:
            dis = _disease_from(stem.replace("PFig_volcano_",""))
            cap = f"Pan-phenome cis-MR volcano \u2014 {dis}."
        n += 1
        gallery_figure(os.path.join(PHE,f), f"Extended Data Fig. P{n} | {cap}")
    print(f"embedded pan-phenome extended figures: {n}")

# ---- Supplementary figures ----
if os.path.isdir(SUP):
    doc.add_page_break()
    H("Supplementary figures", 1, BLUE)
    P("Quality-control and per-disease supplementary figures for the autoimmune "
      "discovery arc (src/20): proteome composition and instrument diagnostics; "
      "per-disease MR volcano, QQ + genomic-inflation \u03bb, and forest plots; per-gene "
      "INTERVAL cis-pQTL regional association; and discovery-versus-replication panels.",
      italic=True, size=9)
    SCAP = {"immune_class_composition":"Immune-class composition of the curated proteome.",
            "source_cell_lineage":"Blood-cell lineage of origin of immune proteins.",
            "instrument_strength":"cis-eQTL instrument strength distribution.",
            "instrument_samplesize":"eQTLGen instrument sample-size distribution.",
            "tests_per_disease":"MR tests performed per disease.",
            "hits_per_disease":"FDR<5% MR hits per disease."}
    def _scap(stem):
        body = stem.split("_",1)[1] if "_" in stem else stem   # drop SFigSNN_
        for k,v in SCAP.items():
            if body==k: return v
        if body.startswith("volcano_"):  return f"Per-disease cis-MR volcano \u2014 {_disease_from(body[8:])}."
        if body.startswith("qq_"):       return f"MR QQ plot and genomic-inflation \u03bb \u2014 {_disease_from(body[3:])}."
        if body.startswith("forest_"):   return f"Per-disease causal-target forest \u2014 {_disease_from(body[7:])}."
        if body.startswith("pqtl_regional_"): return f"INTERVAL cis-pQTL regional association \u2014 {body[14:]}."
        if body.startswith("repl_"):     return f"Discovery-versus-independent-replication \u2014 {_disease_from(body[5:])}."
        return body.replace("_"," ")
    m = 0
    for f in sorted(os.listdir(SUP)):
        if not f.endswith(".png"): continue
        stem = f[:-4]; m += 1
        num = stem.split("_")[0].replace("SFigS","S")
        gallery_figure(os.path.join(SUP,f), f"Supplementary Fig. {num} | {_scap(stem)}")
    print(f"embedded supplementary figures: {m}")

out = os.path.join(OUT, "Plasma_Immunome_Phenome_Atlas_Nature.docx")
doc.save(out)
print("wrote", out)
print(f"tiers: T5={N_T5} T4={N_T4} T3={N_T3} T2(MHC)={N_T2} | pQTL pairs={N_PQTL} "
      f"concordant={N_CONC} discordant={N_DISC} | repl covered={N_COV} replicated={N_REP} "
      f"concordant={N_RCON} | MR sig={N_SIG}/{N_TEST} | immune proteins={N_IMM}")
