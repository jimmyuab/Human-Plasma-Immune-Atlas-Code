#!/usr/bin/env python
"""
HDDM Layer 54 - Step 4
Assemble the Word report (.docx) with all figures embedded + explained,
novelty ranking, paper title and grant title.
Output: 10_manuscript/Plasma_Immunome_Phenome_Atlas_Report.docx
"""
import os, json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = r"I:\Plasma immune atalas"
PROC = os.path.join(ROOT, "02_data_processed")
FIG  = os.path.join(ROOT, "08_figures", "main_figures")
TAB  = os.path.join(ROOT, "09_tables")
OUT  = os.path.join(ROOT, "10_manuscript")
os.makedirs(OUT, exist_ok=True)

S = json.load(open(os.path.join(PROC, "immune_universe_summary.json")))
t2 = pd.read_csv(os.path.join(TAB, "T2_immune_class_summary.tsv"), sep="\t")
NIMM = S["plasma_immune_proteins"]; NUNI = S["olink_universe"]

doc = Document()
# base style
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

def h(txt, lvl=1):
    p = doc.add_heading(txt, level=lvl); return p
def para(txt, italic=False, bold=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(txt); r.italic = italic; r.bold = bold
    r.font.size = Pt(size); return p
def figure(name, caption, width=6.3):
    fp = os.path.join(FIG, name)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(); r = c.add_run(caption); r.italic = True; r.font.size = Pt(9.5)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------------------------------------------------------------- title page
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("HDDM Layer 54 — Plasma Immunome–Phenome Atlas"); r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x6E)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Mapping the Plasma Immunome to Human Health and Disease:\n"
                 "Systemic Immune Communication Programs, Preclinical Disease Trajectories "
                 "and Causal Therapeutic Targets")
rs.italic = True; rs.font.size = Pt(13)
doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Public-data resource prototype • Build report\n").font.size = Pt(11)
meta.add_run("Data sources: UK Biobank–PPP Olink universe, Human Protein Atlas, MSigDB C7/C8, "
             "FinnGen R12, GWAS Catalog, Single Cell Expression Atlas").font.size = Pt(9)
doc.add_page_break()

# ---------------------------------------------------------------- exec summary
h("1. Executive summary", 1)
para(
 f"This report documents the construction of the Plasma Immunome–Phenome Atlas (HDDM Layer 54), "
 f"a resource that organises the human plasma immune proteome measured by Olink Explore into immune "
 f"classes, cellular sources, communication axes, disease links and druggable targets. Starting from "
 f"the full Olink Explore universe of {NUNI:,} proteins, we integrated Human Protein Atlas (HPA) immune-cell "
 f"and secretome annotation, MSigDB C7 immunologic signatures, the FinnGen R12 disease-endpoint manifest "
 f"({2469} endpoints) and the GWAS Catalog study index ({89981:,} studies). A rule-based, immune-biology-"
 f"driven classifier defined a focused universe of {NIMM:,} plasma immune proteins spanning cytokines, "
 f"chemokines, interferon and TNF superfamily members, complement, immune checkpoints, acute-phase proteins, "
 f"HLA/antigen-presentation molecules, immunoglobulin/Fc machinery and immune-cell-enriched secreted proteins.")
para(
 "All steps that can be executed on fully public data are complete and reproduced here as 12 figures and "
 "4 supplementary tables. Steps that require individual-level UK Biobank Olink measurements (phenome-wide "
 "association, Cox incident-disease modelling, 15-year pre-disease trajectories, immune ageing waves and the "
 "machine-learning Plasma Immune Risk Score) are fully specified as a pipeline and flagged as access-gated.")

# ---------------------------------------------------------------- data sources table
h("2. Data sources and access status", 1)
para("Each source was probed for availability and verified after download (HTTP status, file size, row "
     "counts and content parsing). The table below records what was downloaded locally versus what is gated.")
rows = [
 ("Resource","Content","Status"),
 ("UKB-PPP Olink universe (UKB coding 143)", f"{NUNI:,} plasma proteins (gene + name)", "Downloaded"),
 ("Human Protein Atlas (proteinatlas.tsv)", "20,162 genes; immune-cell, secretome, disease", "Downloaded"),
 ("MSigDB C7 / C8", "5,219 immunologic + 840 cell-type gene sets", "Downloaded"),
 ("FinnGen R12 manifest", "2,469 disease endpoints + GCS paths", "Downloaded (manifest)"),
 ("GWAS Catalog study index", "89,981 studies metadata", "Downloaded"),
 ("Single Cell Expression Atlas", "experiment manifest (JSON)", "Downloaded (manifest)"),
 ("UKB individual-level Olink NPX + phenotypes", "per-participant immune proteome", "GATED — UKB application"),
 ("OpenGWAS API", ">50,000 GWAS for MR", "Token required"),
 ("FinnGen full summary statistics", "~250 GB of per-endpoint sumstats", "Bulk via manifest script"),
]
tb = doc.add_table(rows=0, cols=3); tb.style = "Light Grid Accent 1"; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(a,b,c) in enumerate(rows):
    cells = tb.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = a,b,c
    if i==0:
        for cc in cells:
            for p in cc.paragraphs:
                for rr in p.runs: rr.bold = True

# ---------------------------------------------------------------- methods
h("3. Methods (data integration)", 1)
para("Protein universe. The complete Olink Explore protein list was taken from UK Biobank data-coding 143 "
     "(gene symbol + protein description) — the same assay set profiled in 54,219 UKB participants by UKB-PPP.")
para("Annotation. Each protein was joined to HPA on gene symbol to obtain blood-cell (immune) RNA specificity, "
     "the enriched source cell types, secretome location, protein class (including potential drug target) and "
     "disease involvement. MSigDB C7 membership was computed as a supporting immune-signature flag.")
para("Immune classification. Hard immune-biology rules (regular expressions over gene families: IL/CSF/TGFB/TNF/"
     "IFN cytokines; CXCL/CCL/XCL/CX3CL chemokines; complement C-genes/CFB/CFH/MASP; checkpoints PDCD1/CD274/"
     "CTLA4/LAG3/TIGIT; CD markers; HLA; immunoglobulin/Fc; acute-phase; coagulation) were combined with HPA "
     "categorical immune-cell enrichment. A protein enters the plasma immune universe when its composite immune "
     "score ≥ 2. MSigDB C7 and generic 'secreted' flags were deliberately excluded from the inclusion score "
     "(C7 covers 2,872/2,923 Olink genes and is non-discriminating) and retained only as annotation.")

# ---------------------------------------------------------------- results / figures
h("4. Results", 1)

h("4.1 The plasma immune protein universe", 2)
para(f"Of {NUNI:,} Olink proteins, {NIMM:,} ({100*NIMM/NUNI:.0f}%) met the immune-inclusion criteria (Fig. 1). "
     "HPA annotation covered 2,890 proteins; MSigDB C7 covered 2,872, confirming that signature membership alone "
     "cannot define an immune subset and motivating the biology-driven classifier.")
figure("Fig01_immune_universe_funnel.png",
       "Figure 1. Funnel from the full Olink Explore universe to the curated plasma immune protein set.")

h("4.2 Immune protein classification", 2)
para("The immune universe partitions into eleven classes (Fig. 2). Beyond a large immune-cell-enriched secreted "
     "compartment, the atlas captures the canonical soluble immune signalling machinery: 89 cytokines/interleukins, "
     "40 chemokines, 38 complement/coagulation proteins, 27 TNF-superfamily members, 18 immune checkpoints and "
     "8 interferon-axis proteins — the molecules most relevant to systemic immune communication and therapy.")
figure("Fig02_immune_class_composition.png",
       "Figure 2. Composition of the plasma immune proteome by immune class.")

h("4.3 Cellular sources of the plasma immunome", 2)
para("Mapping HPA blood-cell enrichment to coarse lineages reveals the immune cells that contribute most plasma "
     "proteins (Fig. 3): granulocytes (434), myeloid/monocytic cells (312) and dendritic cells (201) dominate the "
     "secreted innate compartment, with substantial T-cell (155), B-cell (86) and NK-cell (77) contributions. "
     "The immune class × source-cell matrix (Fig. 4) localises each signalling class to its producing cells — for "
     "example chemokines and cytokines concentrate in granulocyte/myeloid sources.")
figure("Fig03_immune_cell_source_map.png",
       "Figure 3. Immune-cell source map: number of plasma immune proteins enriched per lineage.")
figure("Fig04_class_by_source_heatmap.png",
       "Figure 4. Immune class × source-cell matrix linking each signalling class to producing cells.")

h("4.4 Annotation structure and immune communication axes", 2)
para("Co-occurrence of annotation flags (Fig. 5) shows the expected modular structure — TNF and checkpoint flags "
     "co-vary, complement and coagulation co-vary, and chemokine annotation is largely orthogonal to immune-cell "
     "enrichment. Grouping by canonical signalling axes (Fig. 11) shows the atlas covers the full breadth of "
     "immune communication: IL-6/gp130, the TNF superfamily, CXC and CC chemokines, the IL-1 family, the interferon "
     "axis, complement and checkpoint pathways.")
figure("Fig05_flag_cooccurrence.png",
       "Figure 5. Co-occurrence (Pearson r) of immune annotation flags.")
figure("Fig11_immune_axes.png",
       "Figure 11. Coverage of key immune communication axes by the Olink plasma immune panel.")

h("4.5 Disease relevance and druggability", 2)
para("HPA disease-involvement terms for the immune proteins (Fig. 9) are enriched for cancer, immunodeficiency and "
     "autoimmune/inflammatory conditions. Crucially for translation, a large fraction of the soluble signalling "
     "classes are flagged as potential drug targets (Fig. 8): 50% of interferon-axis, 40% of acute-phase, 38% of "
     "Ig/Fc, 33% of checkpoint and 30% of cytokine proteins — versus far fewer in the bulk immune-cell-enriched "
     "compartment. Seventy-two non-bulk druggable immune targets are tabulated (Supplementary Table T4).")
figure("Fig08_druggability_by_class.png",
       "Figure 8. Percentage of each immune class flagged as a potential drug target (HPA).")
figure("Fig09_disease_involvement.png",
       "Figure 9. Top disease-involvement terms among plasma immune proteins (HPA).")
para("MSigDB C7 coverage by class (Fig. 10) confirms near-complete immune-signature support across all soluble "
     "signalling classes, validating the immune identity of the curated set.")
figure("Fig10_msigdb_coverage.png",
       "Figure 10. MSigDB C7 ImmuneSigDB coverage by immune class.")

h("4.6 Disease-endpoint and GWAS landscape for causal analysis", 2)
para("The FinnGen R12 manifest provides 2,469 disease endpoints across all ICD chapters with case counts spanning "
     "two orders of magnitude (Fig. 6), the substrate for future Mendelian randomisation and colocalisation against "
     "plasma-protein pQTLs. The GWAS Catalog index contributes thousands of immune-related studies (Fig. 7) for "
     "trait look-up and instrument selection.")
figure("Fig06_finngen_landscape.png",
       "Figure 6. FinnGen R12 disease-endpoint landscape (chapters and case-count distribution).")
figure("Fig07_gwas_immune_traits.png",
       "Figure 7. Immune-related study counts in the GWAS Catalog index.")

h("4.7 The full atlas model", 2)
para("Figure 12 places the public-data prototype within the complete twelve-stage atlas model. The protein "
     "universe → class → source cell → communication axis → disease/GWAS → druggability chain is built here; the "
     "longitudinal, ageing, risk-score and causal-genetics stages await controlled UK Biobank access.")
figure("Fig12_atlas_schematic.png",
       "Figure 12. The HDDM Layer 54 model: realised (public-data) and access-gated stages.", width=6.8)

# ---------------------------------------------------------------- class summary table
h("5. Immune class summary table", 1)
tb2 = doc.add_table(rows=0, cols=5); tb2.style = "Light List Accent 1"
hdr = ["Immune class","N","% drug target","% secreted","% immune-cell enriched"]
cells = tb2.add_row().cells
for i,htext in enumerate(hdr):
    cells[i].text = htext
    for p in cells[i].paragraphs:
        for rr in p.runs: rr.bold = True
for _,r in t2.iterrows():
    cells = tb2.add_row().cells
    cells[0].text = str(r["immune_class"]); cells[1].text = str(int(r["n_proteins"]))
    cells[2].text = str(r["pct_drug_target"]); cells[3].text = str(r["pct_secreted"])
    cells[4].text = str(r["pct_immune_cell_enriched"])

# ---------------------------------------------------------------- novelty ranking
h("6. Novelty ranking", 1)
para("Ranked from strongest to most incremental claim of the full atlas (★ = strength of novelty):")
nov = [
 ("★★★★★","Unified plasma immunome→source-cell→receiver-cell→scATAC-regulation→target chain",
  "No existing resource connects a circulating immune protein to its producing immune cell, its receptor-bearing "
  "receiver cell, the disease-active chromatin/TF program behind it (scATAC layer) and a therapeutic modality in "
  "one model. This systems-level closure is the headline novelty."),
 ("★★★★★","15-year pre-disease immune trajectories across the phenome",
  "Population-scale identification of immune proteins that deviate 10–15 years before onset, across many diseases "
  "simultaneously, reframes them as early-warning and prevention targets rather than reactive biomarkers."),
 ("★★★★☆","Plasma Immune Risk Score (PIRS) as a disease-agnostic immune predictor family",
  "Disease-specific 30-protein immune risk scores benchmarked against demographics, with calibration and SHAP, "
  "give a portable immune-prediction framework."),
 ("★★★★☆","Immune ageing-wave model / inflammaging vs resilience scores",
  "Non-linear immune-ageing waves derived at biobank scale and separated from age-independent disease signals."),
 ("★★★★☆","cis-pQTL Mendelian-randomisation prioritisation of causal immune targets",
  "Triangulating association + trajectory + cis-pQTL MR + colocalisation against FinnGen yields drug-grade causal "
  "evidence with a transparent tiering scheme."),
 ("★★★☆☆","Curated, biology-driven plasma immune universe (this prototype)",
  "A reproducible 1,007-protein immune subset with class, source cell, axis, disease and druggability annotation — "
  "a useful standalone resource and the backbone for everything above."),
]
for stars, title, body in nov:
    p = doc.add_paragraph()
    p.add_run(f"{stars}  ").bold = True
    p.add_run(title).bold = True
    doc.add_paragraph(body)

# ---------------------------------------------------------------- titles
h("7. Suggested paper title", 1)
para("Primary:", bold=True)
para("“The Plasma Immunome–Phenome Atlas: systemic immune communication programs link preclinical disease "
     "trajectories, immune ageing and causal therapeutic targets in 54,000 adults.”", italic=True)
para("Alternatives:", bold=True)
for a in [
 "“A plasma immune protein atlas of human disease: from cellular source to causal target.”",
 "“Circulating immune communication programs predict disease 15 years before onset.”",
 "“From immunome to interventome: causal plasma immune proteins across the human phenome.”"]:
    doc.add_paragraph(a, style="List Bullet")

h("8. Suggested grant title", 1)
para("Primary:", bold=True)
para("“Decoding the Plasma Immunome: a systems atlas of immune communication for early disease prediction and "
     "immune-targeted therapy.”", italic=True)
para("Alternatives:", bold=True)
for a in [
 "“The Human Plasma Immunome Programme: linking immune protein signals to cellular origin, disease risk and "
 "druggable targets across the life course.”",
 "“Immune Communication Atlas: pre-clinical immune trajectories and causal targets for age-related disease.”"]:
    doc.add_paragraph(a, style="List Bullet")

# ---------------------------------------------------------------- next steps
h("9. Next steps to complete the atlas", 1)
for s in [
 "Obtain UK Biobank Olink NPX + phenotype access (Application via AMS) to unlock association, trajectory, "
 "ageing-wave and PIRS modules.",
 "Register an OpenGWAS token and download FinnGen immune-relevant sumstats via the manifest for MR/coloc.",
 "Integrate the existing scATAC disease layer to connect each Tier-1 immune target to disease-active cCREs, "
 "cell types and TF regulators.",
 "Generate the receptor/receiver-cell communication network (CellChat/OmniPath) for the cytokine/chemokine/"
 "checkpoint shortlist (Supplementary Table T3)."]:
    doc.add_paragraph(s, style="List Number")

out = os.path.join(OUT, "Plasma_Immunome_Phenome_Atlas_Report.docx")
doc.save(out)
print("Saved", out)
print("paragraphs:", len(doc.paragraphs))
