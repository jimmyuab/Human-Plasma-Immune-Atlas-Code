#!/usr/bin/env python
"""
HDDM Layer 54 - Step 30
Detailed methodology monograph (~40 pages) for the Plasma Immunome-Phenome Atlas
and the Plasma Immune Risk Score (PIRS) model.

Covers, end to end and in depth:
  - every data source, how it was acquired, and its provenance/licence
  - proteome curation into the plasma immunome
  - genetic instrument construction (cis-eQTL / cis-pQTL)
  - causal inference (Mendelian randomization) and its assumptions
  - statistical colocalization
  - protein-level validation and eQTL-pQTL concordance
  - independent replication
  - the pan-phenome expansion, pleiotropy, cell-source and direction mapping
  - the integrated novelty-priority engine
  - the PIRS supervised model: from data -> features -> algorithm -> score
  - why one would train with it, what questions it answers, and how research benefits
  - evidence tiering, limitations, reproducibility, glossary and references

Output: 10_manuscript/PIRS_and_Atlas_Methodology.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

ROOT = r"I:\Plasma immune atalas"
OUT  = os.path.join(ROOT, "10_manuscript")
os.makedirs(OUT, exist_ok=True)

BLUE  = RGBColor(0x1f, 0x4e, 0x79)
GREY  = RGBColor(0x55, 0x55, 0x55)
DARK  = RGBColor(0x22, 0x22, 0x22)

doc = Document()
base = doc.styles["Normal"]; base.font.name = "Calibri"; base.font.size = Pt(10.5)
base.paragraph_format.space_after = Pt(6)
base.paragraph_format.line_spacing = 1.15

# ---------- helpers ----------
def H1(txt):
    doc.add_page_break()
    p = doc.add_heading(txt, level=1)
    for r in p.runs: r.font.color.rgb = BLUE
    return p
def H1nobreak(txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs: r.font.color.rgb = BLUE
    return p
def H2(txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(0x2e,0x62,0x9e)
    return p
def H3(txt):
    p = doc.add_heading(txt, level=3)
    for r in p.runs: r.font.color.rgb = GREY
    return p
def P(txt, italic=False, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    return p
def BUL(txt, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(txt)
    return p
def NUM(txt, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(txt)
    return p
def EQ(txt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.name = "Consolas"; r.font.size = Pt(10); r.font.color.rgb = DARK
    return p
def CODE(lines):
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln); r.font.name = "Consolas"; r.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(0)
def IMG(relpath, width_in=6.6, caption=None):
    full = os.path.join(ROOT, relpath) if not os.path.isabs(relpath) else relpath
    if not os.path.exists(full):
        return None
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(full, width=Inches(width_in))
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    return p
def TABLE(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.paragraphs[0].add_run(h).bold = True
        for r_ in c.paragraphs[0].runs: r_.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text = str(v)
            for para in cells[i].paragraphs:
                for run in para.runs: run.font.size = Pt(8.5)
    return t

# ================= COVER =================
for _ in range(2): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Plasma Immunome\u2013Phenome Atlas & the Plasma Immune Risk Score (PIRS)")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("A Detailed Methodology: From Raw Data to a Trained Causal-and-Predictive Model")
r.font.size = Pt(14); r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()
for line in [
    "HDDM Layer 54 \u00b7 An open, genetics-anchored resource",
    "Data acquisition \u00b7 proteome curation \u00b7 Mendelian randomization \u00b7 colocalization",
    "protein-level validation \u00b7 independent replication \u00b7 pan-phenome causal mapping",
    "novelty-priority engine \u00b7 supervised Plasma Immune Risk Score",
]:
    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = q.add_run(line); rr.font.size = Pt(11); rr.font.color.rgb = DARK
doc.add_paragraph()
q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = q.add_run("A methods paper \u00b7 built entirely from public data \u00b7 fully reproducible")
rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = GREY
doc.add_paragraph()
q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = q.add_run("HDDM Layer 54 \u00b7 Plasma Immunome\u2013Phenome Atlas Consortium (open resource)")
rr.font.size = Pt(9.5); rr.font.color.rgb = GREY

# ================= ABSTRACT =================
H1nobreak("Abstract")
P("Circulating immune proteins mediate host defence and are the targets of many approved therapies, yet "
  "a systematic, openly reproducible framework linking the plasma immune proteome to disease through "
  "causal genetics \u2014 and a matching predictive model \u2014 has been lacking. We present a complete "
  "methodology for (i) the Plasma Immunome\u2013Phenome Atlas, a genetics-anchored causal map of circulating "
  "immune proteins across 28 diseases in five categories, and (ii) the Plasma Immune Risk Score (PIRS), "
  "a supervised model that converts an individual\u2019s plasma immune proteome into calibrated, "
  "disease-specific risk. Starting from the 2,923-analyte Olink Explore universe, we curate 1,007 "
  "plasma-detectable immune proteins using orthogonal Human Protein Atlas and MSigDB annotation, "
  "instrument 812 immune genes with blood cis-eQTLs (eQTLGen), and estimate causal effects on FinnGen "
  "R12 disease liability by Mendelian randomization. Signals are triangulated by statistical "
  "colocalization (coloc.abf), validated at the protein level with INTERVAL plasma cis-pQTLs, and "
  "replicated in independent, non-FinnGen consortium GWAS. The pipeline blindly recovers approved-drug "
  "axes with the correct pharmacological direction (IL6ST/tocilizumab, CTLA4/abatacept, TNFRSF1A/"
  "anti-TNF), providing an internal positive control, and nominates colocalized targets across "
  "autoimmune, cardiovascular, metabolic, renal and neurodegenerative disease. An integrated "
  "novelty-priority engine ranks targets by cumulative orthogonal evidence while penalising known-drug "
  "and MHC axes, yielding 45 novel colocalized nominations with explicit therapeutic direction. Every "
  "finding is bound to an evidence tier and no claim exceeds its tier. The PIRS layer is "
  "data-source-agnostic and ships no individual-level data: it restricts features to the curated "
  "immunome, fits a cross-validated elastic-net survival model, and returns per-protein weights and "
  "calibrated discrimination in any authorised cohort. All inputs are public and all code is released; "
  "the resource is a reproducible platform for immune target discovery and risk prediction.")
kw = doc.add_paragraph()
r = kw.add_run("Keywords:  "); r.bold = True; r.font.size = Pt(9.5)
kw.add_run("plasma proteome; immune proteins; Mendelian randomization; colocalization; pQTL; "
           "drug target discovery; risk prediction; elastic-net survival; FinnGen; open reproducible resource").font.size = Pt(9.5)

# ================= 0. HOW TO READ =================
H1("1.  Purpose and scope")
P("This monograph is the complete, self-contained methodology behind two linked deliverables: "
  "(i) the Plasma Immunome\u2013Phenome Atlas, a genetics-anchored causal map that connects circulating "
  "immune proteins to human disease, and (ii) the Plasma Immune Risk Score (PIRS), a supervised model "
  "that turns an individual\u2019s plasma immune-protein profile into a calibrated, disease-specific risk "
  "estimate. It is written so that a researcher with no prior exposure to the project can reproduce "
  "every step, understand every modelling choice, and know precisely what each result does and does not "
  "support. The atlas answers the question \u201cwhich immune proteins causally influence which diseases, "
  "and in which direction?\u201d; the PIRS answers the complementary question \u201cgiven a person\u2019s measured "
  "immune proteome today, what is their future risk of a given disease?\u201d The two are deliberately "
  "coupled: the causal atlas tells you which proteins are mechanistically credible, and the PIRS lets "
  "you use them predictively in any cohort you are authorised to analyse.")
P("A guiding principle runs through the entire pipeline and through this document: claim strength is "
  "matched to evidence strength. Association is not causation; a transcript (eQTL) is not a protein "
  "(pQTL); a Mendelian-randomization estimate at the major histocompatibility complex (MHC) is not a "
  "clean instrument; and a cross-validated risk score is not proof of mechanism. Every finding is "
  "therefore reported at an explicit evidence tier, and the language used for it is bound to that tier. "
  "This document explains not only how each number was produced but why each number is allowed to carry "
  "only a certain weight.")
P("The document is organised as a narrative arc. Section 2 defines the model\u2019s scope, its layers and "
  "which parts are plasma-derived. Part A describes the data. Part B describes how the "
  "immune proteome was defined. Part C\u2013F describe the causal-inference engine (instruments, Mendelian "
  "randomization, colocalization, protein-level validation, replication). Part G describes the "
  "pan-phenome expansion and the novelty-priority engine. Part H is the PIRS model in full: data \u2192 "
  "features \u2192 algorithm \u2192 training \u2192 scoring. Part I explains why one would train with the model, what "
  "questions it answers, and how a research programme benefits. Part J covers limitations, "
  "reproducibility, a glossary and references.")

# ================= 2. SCOPE AND FOCUS =================
H1("2.  What this model focuses on (scope, layers and data provenance)")
P("Because the resource joins several kinds of data, it is important to state plainly what it is about, "
  "what it is not about, and \u2014 since the word \u201cplasma\u201d appears throughout \u2014 exactly which parts are "
  "plasma-derived and which are not. This section is written for readers who want to understand the "
  "resource before the technical detail.")

H2("2.1  The core object: circulating plasma immune proteins")
P("The subject of the model is molecular, and specifically it is the plasma (blood-circulating) immune "
  "proteome \u2014 the set of immune-related proteins that are measurable in blood. Concretely this is the "
  "1,007-protein plasma immunome curated from the Olink Explore panel. This is what the model measures, "
  "what the Plasma Immune Risk Score (PIRS) is trained on, and what is ultimately nominated as a "
  "therapeutic target. The model is NOT about the whole proteome, and it is NOT about other molecular "
  "layers such as metabolites, lipids or intracellular/tissue proteins; it is deliberately restricted to "
  "the circulating, secreted immune protein subset.")

H2("2.2  The layers the model connects")
P("Around that molecular core, the resource deliberately links five further layers, so that a plasma "
  "protein is never viewed in isolation but always in relation to its cause, its cellular origin, and "
  "its disease consequence. The model\u2019s focus is best understood as a bridge from molecule to disease:")
TABLE(["Focus / layer","Level","What it contributes to the model"],
      [["Plasma immune protein","molecular \u2014 protein","the CORE: measured features, PIRS inputs, nominated targets"],
       ["Blood gene expression","molecular \u2014 transcript (eQTL)","the causal instrument (why a protein moves disease risk)"],
       ["Germline genetics","molecular \u2014 DNA variants (GWAS)","the anchor that makes findings causal, not merely correlational"],
       ["Cell of origin","cellular","maps each protein to its blood-cell lineage (granulocyte/T/B/NK/myeloid)"],
       ["Disease phenome","phenotypic \u2014 28 diseases","the outcome: what the proteins causally act upon"],
       ["Direction & druggability","translational","turns a target into block-vs-agonise, modality and novelty rank"]])
P("In one sentence: the model focuses on plasma immune proteins (molecular), causally linked through "
  "genetics (transcript and DNA) to disease (phenotypic), mapped to their cell of origin (cellular), and "
  "scored for therapeutic direction and novelty (translational). It is a molecular-to-disease bridge, "
  "not a single-source molecular catalogue.")

H2("2.3  Which parts are plasma, and which are not")
P("The protein and prediction layers are genuinely plasma-derived; the causal-anchor layers are not, and "
  "this is by design \u2014 causal inference needs a genetic anchor, which cannot come from a plasma protein "
  "measurement alone. Stating this honestly prevents the resource from being mistaken for a pure-plasma "
  "dataset.")
TABLE(["Layer","Source","Plasma-derived?","Molecular type"],
      [["Protein universe / PIRS features","Olink Explore","Yes (plasma)","protein"],
       ["Protein-level validation","INTERVAL cis-pQTL","Yes (plasma)","protein"],
       ["Causal instrument","eQTLGen cis-eQTL","No (whole-blood cells)","transcript / RNA"],
       ["Annotation","Human Protein Atlas","No (cells / tissue)","RNA / secretome"],
       ["Disease outcome","FinnGen / OpenGWAS","No (germline DNA + clinical diagnosis)","DNA \u2192 phenotype"]])
P("So the resource is a plasma-immunome model in the sense that its object and its risk score are plasma "
  "proteins, but its causal spine is multi-omic (plasma protein + blood transcript + germline DNA + "
  "clinical phenotype). The transcript-to-protein step exists precisely because the discovery instrument "
  "is transcript rather than plasma protein, and it exposes real discordances (for example IL6ST, whose "
  "effect reverses sign between transcript and plasma protein).")

H2("2.4  What the trained model can answer, in plain terms")
P("A reader deciding whether the resource is useful to them should know the concrete questions it can "
  "answer. At the level of the causal atlas:")
BUL("i.e. which plasma immune proteins causally influence a given "
    "disease, and in which direction (raising or lowering risk).", bold_lead="Which proteins matter for a disease? \u2014 ")
BUL("i.e. is the disease association driven by the same "
    "causal variant as the protein (colocalization), or is it linkage-disequilibrium coincidence?",
    bold_lead="Is the link real or an artefact? \u2014 ")
BUL("i.e. does the signal hold at the protein level (pQTL) and in an "
    "independent, non-Finnish cohort (replication)?", bold_lead="Does it survive validation? \u2014 ")
BUL("i.e. should a drug block/neutralise the protein (it raises risk) or "
    "agonise/replace it (it is protective)?", bold_lead="Which way should a drug push it? \u2014 ")
BUL("i.e. which proteins act across several disease systems, and which "
    "are novel and druggable rather than already-drugged?", bold_lead="Which targets are new and shared? \u2014 ")
P("And at the level of the trained PIRS, on any cohort the user is authorised to analyse:")
BUL("given a person\u2019s measured plasma immune proteome today, what is "
    "their relative future risk of disease X over follow-up?", bold_lead="What is this individual\u2019s risk? \u2014 ")
BUL("which immune proteins carry the predictive signal for disease X in "
    "this cohort, and with what sign?", bold_lead="What drives the prediction? \u2014 ")
BUL("do the proteins the predictor relies on match the proteins the "
    "causal atlas nominates \u2014 i.e. is prediction being driven by mechanism?",
    bold_lead="Is prediction mechanistic? \u2014 ")
BUL("does the immune-protein score add discrimination beyond age, sex "
    "and standard clinical predictors?", bold_lead="Does it add clinical value? \u2014 ")
BUL("does a score trained in one population transfer to another?",
    bold_lead="Does it transport? \u2014 ")
P("Equally important is what it does NOT answer: it is not a diagnostic, it does not establish mechanism "
  "on its own (genetic triangulation nominates, functional experiments prove), it does not cover "
  "non-immune or intracellular proteins, and \u2014 unless the user trains it on their own individual-level "
  "cohort \u2014 the shipped resource is a causal/summary map rather than a per-patient predictor.")

# ================= PART A: DATA =================
H1("PART A \u00b7 DATA: sources, acquisition and provenance")
P("The atlas is built exclusively from public, citable resources. No individual-level, credentialed or "
  "gated dataset was used at any point, and no data were fabricated or imputed to stand in for gated "
  "resources. Where a desirable resource is access-controlled (for example UK Biobank individual-level "
  "proteomics, or the UKB-PPP/Synapse and deCODE plasma pQTL panels), it was explicitly excluded and its "
  "absence is stated as a limitation rather than papered over. This section documents each data source, "
  "how it was obtained, its coordinate system and units, and how it is used downstream.")

H2("A.1  The plasma-protein universe: Olink Explore")
P("The starting universe of measurable proteins is the Olink Explore panel as deployed at population "
  "scale, comprising 2,923 protein assays. Olink Explore uses the Proximity Extension Assay (PEA): each "
  "protein is recognised by a pair of antibodies, each carrying a unique oligonucleotide; when both bind "
  "the same target the oligonucleotides hybridise and are quantified by next-generation sequencing. The "
  "read-out is Normalised Protein eXpression (NPX), a relative, log2 abundance on an arbitrary linear "
  "scale in which one unit corresponds to a doubling of protein concentration. NPX is comparative, not "
  "absolute; it is well-suited to ranking individuals and to regression, which is exactly how the PIRS "
  "uses it. The Olink assay-to-gene mapping (assay identifier, UniProt, gene symbol) is used to bridge "
  "the proteomic space to genetics, because instruments and disease GWAS are indexed by gene and variant, "
  "not by aptamer.")
P("Provenance and access. The identities of the 2,923 assays (the protein universe) are public. The "
  "individual-level NPX measurements themselves, when derived from UK Biobank, are controlled-access and "
  "were NOT used to build the atlas. The PIRS trainer is therefore designed to run on whatever authorised "
  "NPX matrix the end user supplies (UK Biobank under an approved application, a hospital cohort, or a new "
  "Olink run), never on shipped individual-level data.")

H2("A.2  Protein annotation: Human Protein Atlas and MSigDB")
P("Two orthogonal annotation resources convert the flat list of assays into a biologically structured "
  "immune proteome:")
BUL("the master table proteinatlas.tsv (20,162 genes), which provides, per gene, immune-cell RNA "
    "specificity and expression across sorted blood-cell populations, secretome location (secreted / "
    "membrane / intracellular), molecular function, and a curated protein class (including drug-target "
    "and FDA-approved-target flags used later for druggability).", bold_lead="Human Protein Atlas (HPA): ")
BUL("the immunologic-signature collection C7 (5,219 gene sets) and cell-type collection C8 (840 sets). "
    "C7 proved so permissive that it covered almost the entire Olink space; it was therefore demoted to "
    "annotation only and never used as an inclusion filter \u2014 an early, concrete application of the "
    "calibration principle.", bold_lead="MSigDB (Molecular Signatures Database): ")
P("Both resources are public and downloadable in full. HPA supplies the immune-cell lineage of origin "
  "for each protein (granulocyte, monocyte/myeloid, dendritic, T-, B- and NK-cell), which is what allows "
  "later disease signals to be read back to a cell of origin, and the secretome/protein-class fields that "
  "later drive the druggability component of the novelty engine.")

H2("A.3  Genetic instruments (exposure side): eQTLGen cis-eQTLs")
P("Causal direction is exposed by genetics. For the transcript-level exposure we use eQTLGen, a "
  "meta-analysis of blood cis-expression quantitative trait loci (cis-eQTLs) in up to ~31,684 "
  "individuals. The significant-cis file (~322 MB compressed, ~127 million SNP\u2013gene rows) provides, per "
  "variant\u2013gene pair, the Z-score of association between genotype and gene expression, the assessed and "
  "other alleles, and the discovery sample size. Because eQTLGen reports Z-scores rather than betas on a "
  "natural scale, effect sizes are reconstructed analytically (Section C.2). eQTLGen is fully public.")

H2("A.4  Protein-level instruments (validation side): INTERVAL plasma pQTLs")
P("Because a cis-eQTL indexes transcript, not protein, abundance, the atlas seeks an independent "
  "protein-level instrument. We use plasma cis-pQTLs from the INTERVAL study (Sun et al., 2018; ~3,300 "
  "SomaScan aptamers in ~3,300 blood donors), obtained as fully harmonised GWAS-Catalog summary "
  "statistics (GRCh37) over the EBI FTP mirror \u2014 a login-free, public route. For each significant gene "
  "the genome-wide file is downloaded, the cis window (\u00b1500 kb around the gene body, coordinates from "
  "Ensembl GRCh37) is extracted, and the genome-wide file is then deleted to conserve storage. INTERVAL "
  "supplies real effect sizes and standard errors on both the exposure and outcome sides, enabling a "
  "genuine protein-level Mendelian-randomization and protein-level colocalization. The alternative "
  "protein panels (UKB-PPP via Synapse; deCODE via a data-use agreement) are gated behind login or a "
  "legal agreement and were deliberately not used.")

H2("A.5  Disease outcomes: FinnGen R12 (the phenome)")
P("Disease association (the outcome side) comes from FinnGen Data Freeze 12 (R12), a Finnish "
  "biobank-scale GWAS resource covering thousands of endpoints. FinnGen provides per-endpoint genome-wide "
  "summary statistics (variant, alleles, beta, standard error, p-value, allele frequency) with a public "
  "manifest of direct download paths. The atlas uses FinnGen at two scales:")
BUL("13 immune-mediated endpoints (rheumatoid arthritis, multiple sclerosis, ankylosing spondylitis, "
    "psoriasis, coeliac disease, Crohn\u2019s disease, sarcoidosis, systemic lupus, vitiligo, Sj\u00f6gren "
    "syndrome, autoimmune thyroiditis, autoimmune hyperthyroidism, Guillain\u2013Barr\u00e9 syndrome).",
    bold_lead="Discovery (autoimmune core): ")
BUL("15 additional well-powered non-immune endpoints spanning cardiovascular (hypertension, coronary "
    "heart disease, atrial fibrillation, heart failure, stroke, venous thromboembolism), metabolic "
    "(type-1 and type-2 diabetes, obesity), renal (chronic kidney disease), and neurodegenerative/aging "
    "(dementia, Alzheimer\u2019s disease, epilepsy, glaucoma, osteoporosis) categories \u2014 giving 28 diseases "
    "in five categories in total.", bold_lead="Pan-phenome expansion: ")
P("Each FinnGen sumstat file is ~0.8 GB; all 28 were downloaded into the project (~23 GB) so that every "
  "result is computed on real, complete summary statistics rather than a subset. FinnGen R12 is public.")

H2("A.6  Independent replication GWAS: OpenGWAS consortia")
P("To guard against discovery being an artefact of a single (Finnish) cohort, each significant "
  "instrument variant is re-queried in an independent, FinnGen-free consortium GWAS via the OpenGWAS API. "
  "The independent studies used are IMSGC/Patsopoulos (multiple sclerosis), Okada (rheumatoid arthritis), "
  "IGAS/Cortes (ankylosing spondylitis), Stuart (psoriasis), Fischer (sarcoidosis), Sakaue (autoimmune "
  "thyroid) and others. OpenGWAS requires a free personal JWT token; the token is stored locally and "
  "never committed, and because the replication output is saved to disk, the token is not needed to "
  "reproduce downstream steps.")

H2("A.7  Data-access reality and what was deliberately excluded")
P("The single most important honesty constraint of the project is that several attractive resources are "
  "gated, and none of them were substituted with fabricated data. Individual-level UK Biobank NPX and "
  "phenotypes require an approved application through the UKB Access Management System (Olink is field "
  "30900). UKB-PPP plasma pQTLs are served only behind a Synapse login. deCODE plasma pQTLs require "
  "submitting a data-use agreement under the user\u2019s identity. Each of these was left out by design. The "
  "consequence \u2014 that the atlas is a summary-statistics causal resource rather than an "
  "individual-level-trained model \u2014 is stated plainly, and the PIRS is engineered precisely to close "
  "that gap for any user who does hold authorised individual-level data.")

TABLE(["Layer","Source","Role","Access"],
      [["Protein universe","Olink Explore (2,923 assays)","defines measurable proteome","public (identities)"],
       ["Annotation","Human Protein Atlas","immune-cell / secretome / druggability","public"],
       ["Annotation","MSigDB C7/C8","immune signatures (annotation only)","public"],
       ["eQTL exposure","eQTLGen cis-eQTL","transcript-level instruments","public"],
       ["pQTL validation","INTERVAL (Sun 2018)","protein-level instruments","public (GWAS Catalog)"],
       ["Disease GWAS","FinnGen R12 (28 endpoints)","outcome / phenome","public"],
       ["Replication","OpenGWAS consortia","out-of-sample validation","free token"],
       ["Excluded","UKB-PPP / deCODE / UKB NPX","(not used)","gated"]])

# ================= PART B: CURATION =================
H1("PART B \u00b7 Defining the plasma immunome")
P("The purpose of Part B is to convert the 2,923 raw Olink assays into a defensible, biologically "
  "coherent set of plasma-detectable immune proteins \u2014 the feature space shared by every downstream "
  "analysis and by the PIRS.")

H2("B.1  Inclusion score")
P("Each assay is scored by an additive, transparent inclusion rule that combines orthogonal evidence "
  "that a protein is (a) immune-relevant and (b) plausibly present and measurable in plasma. The score "
  "weights most heavily membership of core immune pathways (cytokine, chemokine, interferon, TNF "
  "superfamily), complement and checkpoint biology, and HPA immune-cell RNA enrichment; it additionally "
  "credits secreted/receptor secretome location (which favours plasma detectability and druggability) and "
  "curated CD/leukocyte, HLA, immunoglobulin, acute-phase and coagulation annotation. Proteins passing "
  "the threshold constitute the plasma immunome. The procedure retains 1,007 plasma immune proteins from "
  "the 2,923-assay universe. Crucially, MSigDB C7 was excluded from the score because it covered almost "
  "the entire panel and therefore carried no discriminating information \u2014 an explicit guard against a "
  "criterion that would have admitted everything.")

H2("B.2  Cell-source and druggability annotation")
P("For every retained protein the resource stores its immune class (the pathway that defines it) and its "
  "blood-cell lineage of origin from HPA sorted-cell expression. This is what later permits a disease "
  "signal to be mapped back to granulocyte, myeloid/monocyte, dendritic, T-, B- or NK-cell biology. It "
  "also stores the HPA protein class and drug-target/FDA-target flags, secretome location and membrane "
  "topology, which feed the druggability component of the novelty engine (secreted and single-pass "
  "receptor proteins are the classes most tractable to antibodies and biologics).")

# ================= PART C: INSTRUMENTS =================
H1("PART C \u00b7 Genetic instruments")
P("Mendelian randomization (MR) uses a genetic variant as a natural experiment: because alleles are "
  "assigned at conception and are largely independent of later confounders, a variant that raises a "
  "protein\u2019s abundance and is also associated with disease provides directional, confounding-resistant "
  "evidence that the protein influences the disease. The quality of an MR analysis is the quality of its "
  "instruments; Part C documents how they are built.")

H2("C.1  cis-instrument selection")
P("For each of the 812 immune genes with a usable signal, the single strongest cis-eQTL (the variant "
  "with the largest absolute expression Z-score within the gene\u2019s cis window in eQTLGen) is chosen as a "
  "one-SNP instrument. A cis instrument is preferred to a trans instrument because a variant acting in cis "
  "on the very gene it neighbours is far more likely to affect disease through that gene\u2019s protein "
  "product (satisfying the MR exclusion-restriction assumption) than a distant trans variant, which may "
  "act through many genes. Single strongest-cis instruments are the standard, conservative choice for "
  "protein/expression MR and are robust to the horizontal-pleiotropy problems that afflict multi-variant "
  "polygenic instruments.")

H2("C.2  Reconstructing effect sizes from Z-scores (Zhu et al., 2016)")
P("eQTLGen reports a Z-score, not a beta on an interpretable scale. Using the effect-allele frequency p "
  "and the study sample size n, the standardised effect size and its standard error are reconstructed "
  "with the Zhu et al. (2016) transformation:")
EQ("denom = 2 \u00b7 p \u00b7 (1 \u2212 p) \u00b7 (n + Z\u00b2)")
EQ("beta = Z / \u221adenom          se = 1 / \u221adenom")
P("This yields a beta and standard error on the standardised-genotype scale that are directly usable in "
  "the Wald-ratio MR and, later, in the colocalization variance terms. The same allele frequency (taken "
  "from the FinnGen control/reference side) is used consistently so that exposure and outcome are on "
  "compatible scales.")

H2("C.3  Harmonisation")
P("Exposure and outcome are harmonised to a common effect allele before any ratio is formed. Palindromic "
  "variants (A/T or C/G, whose strand is ambiguous) are removed rather than guessed. Alleles are aligned "
  "so that the eQTL effect is expressed with respect to the expression-increasing allele; the disease "
  "effect is flipped to match. Variants absent from the outcome GWAS are dropped. This harmonisation is "
  "applied identically in the transcript arm, the protein arm and the replication arm so that the "
  "direction of every reported effect is mutually comparable.")

# ================= PART D: MR =================
H1("PART D \u00b7 Causal inference by Mendelian randomization")

H2("D.1  The Wald ratio")
P("With a single instrument, the causal effect of the exposure (protein/transcript abundance) on the "
  "outcome (disease liability) is the Wald ratio \u2014 the ratio of the variant\u2019s effect on the outcome to "
  "its effect on the exposure:")
EQ("beta_MR = beta_outcome / beta_exposure")
EQ("se_MR   = se_outcome / |beta_exposure|      (leading-order / delta method)")
P("The estimate is exponentiated to an odds ratio (OR) so that OR>1 means genetically higher protein "
  "raises disease risk, and OR<1 means it lowers risk. This single directional number is the atomic unit "
  "of the atlas: it is not merely a correlation but a genetically-anchored estimate of the sign and "
  "approximate magnitude of a protein\u2019s causal effect on a disease.")

H2("D.2  Multiple-testing control")
P("Across the discovery arm, 812 immune instruments are tested against 13 immune diseases, giving 8,749 "
  "gene\u2013disease tests. The false-discovery rate is controlled at 5% by the Benjamini\u2013Hochberg "
  "procedure, yielding 32 significant gene\u2013disease pairs. In the pan-phenome arm the same instruments "
  "are tested against all 28 diseases (18,844 tests), yielding 176 FDR-significant pairs. FDR rather than "
  "Bonferroni is used because the tests are correlated (shared instruments, related diseases) and the "
  "aim is discovery of a credible candidate set for downstream triangulation, not a single confirmatory "
  "decision.")

H2("D.3  Assumptions, and how each is addressed")
NUM("the instrument is associated with the exposure. Guaranteed by construction \u2014 the strongest "
    "genome-wide-significant cis-eQTL is chosen.", bold_lead="Relevance: ")
NUM("the instrument shares no confounder with the outcome. Made plausible by using cis variants and by "
    "flagging/withholding the MHC, the archetypal confounded region.", bold_lead="Independence: ")
NUM("the instrument affects the outcome only through the exposure. This is the hard assumption; a "
    "single cis-eQTL can still act through a neighbouring gene or through the protein rather than the "
    "transcript. This is exactly why the atlas does not stop at MR but triangulates with colocalization "
    "(Part E), an orthogonal protein-level instrument (Part F) and independent replication (Part G).",
    bold_lead="Exclusion restriction: ")
P("Two additional guards are built in. First, the MHC region on chromosome 6 (and complement genes such "
  "as C2/C4A/C4B) is flagged for extreme long-range linkage disequilibrium; signals there are capped at "
  "\u201cnomination\u201d and never promoted to a causal claim, because LD there routinely manufactures spurious "
  "colocalization. Second, the positive controls \u2014 genes that are the targets of approved drugs \u2014 are "
  "used as an internal calibration: a pipeline that blindly recovers IL6ST (tocilizumab) in rheumatoid "
  "arthritis, CTLA4 (abatacept) in autoimmune thyroid disease and TNFRSF1A (etanercept) in ankylosing "
  "spondylitis, with the correct direction, is calibrated to be believed when it points somewhere new.")

H2("D.4  Directionality, sensitivity and statistical power")
P("Three further considerations govern whether a Wald-ratio estimate deserves attention. First, "
  "directionality: because both the exposure (expression) and the outcome (disease) are associated with "
  "the same variant, one must be satisfied that the variant acts on expression first and disease second, "
  "not the reverse. Using a cis-eQTL as the instrument makes the expression-first ordering "
  "mechanistically compelling (a nearby regulatory variant plausibly sets transcript level, which is "
  "upstream of a complex disease), and a Steiger-type check \u2014 the variant explains far more variance in "
  "cis-expression than in disease liability \u2014 is consistent with that ordering for the reported hits. "
  "Second, sensitivity: a single-instrument Wald ratio cannot run the multi-instrument pleiotropy "
  "diagnostics (MR-Egger intercept, weighted median, MR-PRESSO) that require many variants, so the atlas "
  "deliberately does not claim to have excluded horizontal pleiotropy by MR alone. Instead it treats "
  "colocalization, protein-level triangulation and out-of-sample replication as the sensitivity analyses, "
  "each of which attacks a different failure mode (LD confounding, transcript-vs-protein discordance, and "
  "cohort-specific artefact respectively). Third, power: the precision of a Wald ratio is bounded by the "
  "strength of the instrument (the cis-eQTL F-statistic) and by the number of disease cases; weak "
  "instruments inflate the ratio\u2019s variance. Because only genome-wide-significant cis-eQTLs are used, "
  "weak-instrument bias is minimised, but the corollary is that genes without a strong cis-eQTL simply "
  "cannot be tested \u2014 an absence of a hit is never interpreted as evidence of no effect.")

# ================= PART E: COLOC =================
H1("PART E \u00b7 Statistical colocalization")
P("MR can be fooled when the exposure variant and the true disease variant are merely in linkage "
  "disequilibrium rather than being one and the same. Colocalization tests exactly this: does the "
  "association signal for the protein and for the disease arise from a single shared causal variant?")

H2("E.1  coloc.abf and the Wakefield approximate Bayes factor")
P("The Giambartolomei et al. (2014) coloc.abf framework evaluates five mutually exclusive hypotheses at "
  "a locus (H0 no signal; H1 exposure only; H2 outcome only; H3 two distinct causal variants; H4 one "
  "shared causal variant) and returns their posterior probabilities. Per SNP, the Wakefield approximate "
  "Bayes factor is computed in log space:")
EQ("r = W / (W + V)        lABF = 0.5 \u00b7 ( log(1 \u2212 r) + r \u00b7 Z\u00b2 )")
P("where V is the variant\u2019s variance (the squared standard error) and W is the prior variance on the "
  "effect (W = 0.15\u00b2 for the eQTL side, 0.20\u00b2 for the GWAS side). Real FinnGen standard errors are used "
  "for the outcome, and Zhu-reconstructed standard errors for the exposure, so the Bayes factors reflect "
  "genuine per-variant precision. The H3 (two-variant) term is evaluated in log space to avoid numerical "
  "underflow. Default priors p1 = p2 = 1\u00d710\u207b\u2074 (a variant is causal for exposure, or for outcome) and "
  "p12 = 1\u00d710\u207b\u2075 (causal for both) are used throughout.")

H2("E.2  Caching the eQTL extract")
P("The 127-million-row eQTLGen file is scanned once and the per-gene cis extracts are cached to a "
  "pickle. Colocalization then runs against the cache, and the cache is extended on demand when the "
  "pan-phenome arm introduces genes not previously seen. This makes the difference between an analysis "
  "that runs in minutes and one that re-reads 127 million rows per gene.")

H2("E.3  Interpretation and promotion rule")
P("A posterior probability of a shared causal variant PP.H4 \u2265 0.8 promotes a locus from a "
  "genetically-supported nomination to a prioritised, transcript-level causal target. The method behaves "
  "as it should on the positive controls (IL6ST\u2013rheumatoid arthritis PP.H4 = 1.00; CTLA4 > 0.98; "
  "TNFSF14\u2013multiple sclerosis > 0.99). In the autoimmune core, 17 loci reach PP.H4 \u2265 0.8; in the "
  "pan-phenome arm, 52 loci do. Loci with strong MR but weak colocalization remain nominations, and MHC "
  "loci are never promoted regardless of PP.H4.")

H2("E.4  The five hypotheses, and what each rejection means")
P("coloc.abf partitions the posterior over five hypotheses, and the value of the method lies in "
  "distinguishing the two that look identical to MR:")
NUM("no association with either trait at the locus.", bold_lead="H0: ")
NUM("a causal variant for the exposure only (the protein signal is real, the disease signal is not here).", bold_lead="H1: ")
NUM("a causal variant for the outcome only (the disease signal is real, the protein signal is not here).", bold_lead="H2: ")
NUM("two DISTINCT causal variants in LD \u2014 the exposure and disease signals are driven by different "
    "variants that merely sit close together. This is the scenario that fools MR into a false positive.",
    bold_lead="H3: ")
NUM("one SHARED causal variant drives both \u2014 the evidence the atlas requires for a causal claim.", bold_lead="H4: ")
P("A high PP.H4 rules out H3 specifically, which is why colocalization is the necessary complement to MR: "
  "MR establishes a directional dose\u2013response, and colocalization establishes that the same variant is "
  "responsible for both ends of it. A locus with strong MR but high PP.H3 is downgraded, because its MR "
  "estimate is most likely an LD artefact.")

H2("E.5  Prior sensitivity and robustness")
P("The posterior depends on the priors p1, p2 and p12. The atlas uses the community-standard defaults "
  "(p1 = p2 = 1\u00d710\u207b\u2074, p12 = 1\u00d710\u207b\u2075), under which the prior odds favour distinct signals over a shared "
  "one; a locus must therefore present strong data to reach PP.H4 \u2265 0.8, making the threshold "
  "conservative. Because the reported colocalizations are typically decisive (positive controls at "
  "PP.H4 \u2248 1.00 and leading nominations at PP.H4 > 0.97), they are insensitive to reasonable "
  "perturbation of p12; a signal that flips across an order-of-magnitude change in the prior would not "
  "have been promoted. The number of shared variants (nsnp) is reported per locus so that thin overlaps "
  "\u2014 where the posterior is data-starved \u2014 can be identified and discounted.")

# ================= PART F: PQTL =================
H1("PART F \u00b7 Protein-level validation and eQTL\u2013pQTL concordance")
P("A cis-eQTL tells you about messenger RNA; a therapeutic acts on protein. The atlas therefore repeats "
  "the entire MR-plus-colocalization logic using the INTERVAL plasma cis-pQTL as a protein-level "
  "instrument for every significant gene that has a matched aptamer.")

H2("F.1  Protein-level MR and colocalization")
P("For each significant gene, the strongest cis-pQTL within \u00b1500 kb drives a Wald-ratio MR against the "
  "disease using real betas and standard errors on both sides, and coloc.abf is applied over the shared "
  "cis variants. Of the significant genes, 18 had a matched INTERVAL aptamer (intracellular and "
  "MHC-region proteins such as CTLA4, C2 and HNMT have no plasma aptamer and cannot be tested this way).")

H2("F.2  The promotion rule and the value of discordance")
P("A gene\u2013disease pair is promoted to a protein-level causal target (the highest tier reached) only if "
  "the transcript- and protein-level MR are directionally concordant AND the protein-level colocalization "
  "reaches PP.H4 \u2265 0.8. Two pairs meet this bar: TNFSF14 in multiple sclerosis and SWAP70 in rheumatoid "
  "arthritis, both colocalizing at transcript and protein (PP.H4 > 0.97). Of the 18 testable pairs, 9 are "
  "directionally concordant and 9 discordant. The discordances are reported, not hidden, because they are "
  "informative: IL6ST reverses sign between transcript and plasma protein, consistent with the known "
  "biology of soluble gp130 as an inhibitor of IL-6 trans-signalling. A discordance tells a drug-hunter "
  "that a transcript-level nomination should not yet be read as a protein-level causal claim \u2014 exactly "
  "the distinction a therapeutic programme needs before committing to a modality.")

# ================= PART G: PHENOME + NOVELTY =================
H1("PART G \u00b7 Pan-phenome expansion and the novelty-priority engine")

H2("G.1  From autoimmunity to a five-category phenome")
P("The same 812 immune instruments are re-run against all 28 diseases, annotated by category "
  "(autoimmune, cardiovascular, metabolic, renal, neuro/aging). This yields 176 FDR-significant causal "
  "gene\u2013disease pairs \u2014 cardiovascular 67, autoimmune 42, metabolic 35, neuro/aging 31, renal 1 \u2014 "
  "turning a single-disease-family resource into a multi-system causal map. New, non-autoimmune signals "
  "emerge (for example ACE and IFNGR2 in hypertension, PSRC1 and PLAUR in coronary heart disease, PROCR "
  "in venous thromboembolism, and ACE colocalizing with dementia and Alzheimer\u2019s disease at PP.H4 > "
  "0.97).")

H2("G.2  Cross-category pleiotropy, cell-source and direction mapping")
P("Three analyses exploit the breadth. (i) Pleiotropy: genes causal in \u2265 2 disease categories are "
  "extracted, exposing shared immune control points that link immune and cardiometabolic disease "
  "(IFNGR2, SWAP70, MERTK, MPO, SPINK8, PIK3IP1; PM20D1 spans three categories). (ii) Cell-source "
  "enrichment: causal targets are mapped back to their blood-cell lineage and tested by Fisher\u2019s exact "
  "test; granulocyte-sourced proteins carry the strongest causal enrichment across the phenome. (iii) "
  "Direction map: each of the 176 signals is encoded as protein-raising-risk (OR>1 \u2192 blockade/"
  "neutralisation strategy) or protein-lowering-protective (OR<1 \u2192 agonism/replacement strategy), with "
  "MHC/LD-confounded loci held separately.")

H2("G.3  The integrated novelty-priority engine")
P("All real evidence layers are combined into a single, auditable priority score per gene\u2013disease pair, "
  "designed to surface genuinely new biology rather than re-discover approved drugs:")
EQ("NoveltyPriority = s_causal + s_coloc + s_pleio + s_drug + s_cell \u2212 p_known \u2212 p_mhc")
P("with components: s_causal = min(\u2212log10 FDR / 5, 2) rewards causal strength; s_coloc = PP.H4 rewards a "
  "shared variant; s_pleio = (n_categories \u2212 1) \u00d7 0.5 rewards cross-system action; s_drug scales with "
  "HPA-derived druggability (FDA-target \u00d72, plus secreted/membrane tractability); s_cell = 0.5 credits a "
  "defined cell source; and the penalties p_known = 1.5 (an approved-drug axis such as CTLA4, IL6ST, "
  "IL2RA) and p_mhc = 2.0 (MHC/LD caution) push known and confounded axes down. Applying the engine "
  "leaves 45 novel colocalized targets, topped by ACE\u2192dementia, IFNGR2\u2192hypertension/psoriasis, "
  "ERBB3\u2192type-1 diabetes, PLAUR\u2192coronary heart disease and SPINK8\u2192type-2 diabetes.")

# ================= PART H: PIRS MODEL =================
H1("PART H \u00b7 The Plasma Immune Risk Score (PIRS): from data to trained algorithm")
P("Everything above is a summary-statistics causal resource. The PIRS is its predictive counterpart: a "
  "supervised model that takes an individual\u2019s measured plasma immune proteome and returns a calibrated, "
  "disease-specific future-risk estimate. Part H documents the model in full \u2014 inputs, feature space, "
  "algorithm, training procedure, outputs, and how to score new individuals \u2014 exactly as implemented in "
  "the trainer (src/22_train_pirs.py).")

H2("H.1  Design philosophy: bring your own data, never fabricate")
P("The trainer ships NO individual-level data and never invents any. It is data-source-agnostic: it runs "
  "on any cohort the user is authorised to analyse (UK Biobank under an approved application, a hospital "
  "cohort, or a fresh Olink run). If input files are absent it prints the exact schema it needs, can "
  "write blank templates, and exits cleanly with status 0 rather than fabricating a result. This is the "
  "same no-fabrication discipline that governs the whole atlas, applied to the predictive layer.")

H2("H.2  Inputs and schema")
P("Three tab- or comma-separated tables are supplied on the command line; the identifier column name is "
  "configurable (default 'id').")
BUL("participant \u00d7 protein matrix. Rows are participants; one column is the id; the remaining columns "
    "are per-protein values (Olink NPX log2, or any normalised expression). Protein columns may be named "
    "by Olink assay ID OR by gene symbol \u2014 both are matched against the curated immune panel.",
    bold_lead="npx: ")
BUL("survival labels, one row per id \u00d7 disease, with columns id, disease, event (1 = incident case, "
    "0 = censored) and time_years (follow-up from baseline to first event or censoring).",
    bold_lead="outcomes: ")
BUL("confounders (id, age, sex, \u2026), merged when present.", bold_lead="covariates (optional): ")

H2("H.3  Feature space: restriction to the curated immunome")
P("The model does not use all measured proteins. It intersects the supplied proteomic columns with the "
  "1,007 curated plasma immune proteins (matching by Olink ID or gene symbol) and trains only on that "
  "immune feature set. This is a deliberate, mechanistically-motivated dimensionality reduction: it "
  "focuses the model on the biology the atlas has causally characterised, improves the "
  "events-per-feature ratio, and makes the resulting weights interpretable as immune biology rather than "
  "as an opaque proteome-wide signature. If none of the supplied columns match the immune panel the "
  "trainer stops with an explicit error rather than silently training on the wrong features.")

H2("H.4  Algorithm: cross-validated elastic-net survival")
P("The core estimator is a penalised (elastic-net) survival model. Elastic net combines an L1 (lasso) "
  "penalty, which performs feature selection by driving most protein weights to exactly zero, with an L2 "
  "(ridge) penalty, which stabilises the estimates when proteins are correlated (as immune proteins "
  "frequently are, moving together within pathways). The mixing parameter l1_ratio (default 0.5) balances "
  "the two. Penalisation is essential here because the immune feature space (up to ~1,007 proteins) is "
  "large relative to the number of incident events in most cohorts; an unpenalised model would overfit.")
P("The trainer auto-selects the best available survival backend so it runs in any environment:")
NUM("scikit-survival CoxnetSurvivalAnalysis \u2014 a true elastic-net Cox proportional-hazards model, the "
    "preferred backend; risk is the linear predictor and discrimination is Harrell\u2019s C-index for "
    "right-censored data.", bold_lead="coxnet: ")
NUM("lifelines CoxPHFitter with an elastic-net penalty \u2014 an equivalent Cox model used when "
    "scikit-survival is unavailable.", bold_lead="lifelines: ")
NUM("an elastic-net logistic regression on a fixed-horizon label (event within the horizon, default 10 "
    "years), scored by AUROC \u2014 a graceful fallback that needs no survival library. It answers a slightly "
    "different question (risk by a fixed horizon rather than the full hazard) but keeps the pipeline "
    "runnable everywhere.", bold_lead="logistic: ")
P("Every backend is wrapped in the same preprocessing pipeline: median imputation of missing protein "
  "values followed by standardisation (zero mean, unit variance). Standardisation puts all proteins on a "
  "common scale so the penalty is applied fairly and so the learned weights are directly comparable "
  "across proteins.")

H2("H.5  Training procedure and honest evaluation")
P("Training proceeds per disease. For a given disease the proteomic matrix is merged with that disease\u2019s "
  "survival labels; diseases with fewer than a minimum number of incident events (default 50) are skipped "
  "with a message, because too few events cannot support a stable high-dimensional model. Discrimination "
  "is estimated by stratified k-fold cross-validation (default 5 folds, stratified on event status): in "
  "each fold the preprocessing is fitted on the training partition ONLY and then applied to the held-out "
  "partition, so no information from the test fold leaks into imputation or scaling. The per-fold C-index "
  "(or AUROC) values are averaged to give an out-of-sample discrimination estimate with a standard "
  "deviation. Only after cross-validated evaluation is the final model refitted on all participants to "
  "produce the deployable weights. In an end-to-end synthetic test the pipeline achieved a cross-"
  "validated C-index of 0.777 and correctly recovered the planted risk proteins, confirming the "
  "estimator, the leakage-free cross-validation and the weight extraction all behave as intended.")

H2("H.6  Outputs")
BUL("per-protein PIRS coefficients for each disease \u2014 the score itself, sorted "
    "by absolute weight, with both the protein identifier and its gene symbol.",
    bold_lead="pirs_<disease>_weights.tsv: ")
BUL("per-disease cross-validated discrimination (C-index or AUROC, mean \u00b1 SD, "
    "event and participant counts, backend used).", bold_lead="pirs_cv_metrics.tsv: ")
BUL("the fitted pipeline (imputer + scaler + model + feature list + backend), "
    "pickled for scoring new individuals.", bold_lead="pirs_<disease>_model.pkl: ")
BUL("a horizontal-bar figure of discrimination across all modelled diseases, "
    "with the 0.5 no-discrimination reference line.", bold_lead="pirs_performance.png: ")

H2("H.7  Scoring a new individual")
P("To score a new person for a given disease: load the corresponding model pickle, take that person\u2019s "
  "values for the model\u2019s feature proteins (in the stored order), apply the saved imputer and scaler, and "
  "evaluate the model. The output is a relative risk/hazard for survival backends or a horizon-risk "
  "probability for the logistic backend. Because the imputer and scaler are stored inside the pickle, a "
  "new individual is transformed with exactly the statistics learned during training, which is what makes "
  "scores comparable across people and across time.")

H2("H.8  Hyperparameters and their tuning")
P("The estimator exposes a small, interpretable set of hyperparameters. The elastic-net mixing parameter "
  "l1_ratio (default 0.5) trades sparsity against stability: values near 1 select a small number of "
  "proteins (useful when a compact, assayable panel is the goal), values near 0 retain more correlated "
  "proteins (useful when the aim is maximal discrimination). The regularisation strength is set along a "
  "coordinate-descent path (Coxnet\u2019s alpha grid with alpha_min_ratio = 0.01), and the deployed "
  "coefficients are taken at the most permissive end of the path after cross-validated evaluation; a user "
  "who wishes to tune the penalty explicitly can nest an inner cross-validation over the alpha grid. The "
  "number of folds (default 5) and the minimum-events threshold (default 50) trade bias against variance "
  "and stability: more folds give a lower-bias but higher-variance performance estimate, and a higher "
  "event floor refuses to train scores that the data cannot support. All are single command-line flags, "
  "so the entire trade-off surface is reproducible and auditable.")

H2("H.9  Calibration versus discrimination")
P("Cross-validated C-index (or AUROC) measures discrimination \u2014 whether higher-risk individuals are "
  "ranked above lower-risk ones \u2014 which is the property most relevant to enrichment and stratification. "
  "It does not by itself guarantee calibration (whether predicted absolute risks match observed "
  "frequencies). For applications that need absolute-risk estimates rather than rankings, the linear "
  "predictor from the Cox model should be combined with an estimated baseline hazard (Breslow) and "
  "assessed with a calibration plot and a calibration-in-the-large check in the target cohort. The "
  "trainer intentionally reports discrimination as the primary, cohort-portable metric and leaves "
  "absolute-risk calibration to the deployment cohort, because calibration is population-specific in a "
  "way that discrimination is less so.")

H2("H.10  Recommended external-validation protocol")
P("A PIRS should be validated before use, and the resource makes the protocol explicit:")
NUM("train and cross-validate in the development cohort; record discrimination and the non-zero "
    "protein weights.", bold_lead="Internal: ")
NUM("apply the frozen model (imputer + scaler + coefficients) to a second, independent cohort with the "
    "same disease definition and follow-up structure; report discrimination and a calibration plot with "
    "no refitting.", bold_lead="Temporal/geographic external: ")
NUM("compare the non-zero PIRS proteins against the atlas\u2019s causal nominations for the same and related "
    "diseases; concordance is evidence that prediction is mechanism-driven and therefore more likely to "
    "transport.", bold_lead="Mechanistic cross-check: ")
NUM("test whether the PIRS adds discrimination beyond age, sex and standard clinical predictors, since "
    "only incremental value justifies a new assay.", bold_lead="Incremental value: ")

H2("H.11  Missing data, batch effects and class imbalance")
P("Three practical realities are handled explicitly. Missing protein values are median-imputed within "
  "the training fold only (never using test-fold information), which is appropriate for the "
  "missing-at-random pattern typical of below-detection Olink readings; users with substantial "
  "informative missingness should consider an indicator-augmented imputation. Batch and plate effects, "
  "which are common in large proteomic runs, should be regressed out or included as covariates before "
  "training, because the score assumes NPX is comparable across participants. Class imbalance \u2014 few "
  "incident cases relative to controls \u2014 is addressed structurally: the survival backends use all "
  "censored information rather than a rare-positive binary label, the cross-validation is stratified on "
  "event status so every fold contains cases, and the minimum-events floor prevents training a score that "
  "would be dominated by a handful of events.")

# ================= PART I: WHY / QUESTIONS / BENEFIT =================
H1("PART I \u00b7 Why train with it, what it answers, and how research benefits")

H2("I.1  Why train a PIRS at all")
P("Single proteins are noisy and pleiotropic; a weighted combination of mechanistically-selected immune "
  "proteins is more stable and more informative than any one marker. Training a PIRS converts a static "
  "proteomic measurement into an actionable, disease-specific risk estimate, and does so on the user\u2019s "
  "own population so that the score is calibrated to their assay, ancestry and clinical setting rather "
  "than borrowed from an external cohort. Because the feature space is restricted to the causally-"
  "characterised immunome, the resulting weights are not a black box: each non-zero coefficient names an "
  "immune protein whose causal credibility for related diseases can be looked up directly in the atlas. "
  "Training is therefore valuable both predictively (a risk tool) and mechanistically (a hypothesis "
  "generator whose features are pre-vetted).")

H2("I.2  Questions the model answers")
BUL("given a person\u2019s plasma immune proteome today, what is their "
    "relative risk of developing disease X over follow-up?", bold_lead="Individual risk: ")
BUL("which immune proteins carry the predictive signal for disease X in this "
    "cohort, and with what sign?", bold_lead="Feature importance: ")
BUL("do the proteins the predictive model relies on coincide with the "
    "proteins the causal atlas nominates \u2014 i.e. is prediction being driven by mechanism?",
    bold_lead="Causal\u2013predictive concordance: ")
BUL("does an immune-protein score add discrimination over age/sex/clinical "
    "covariates for a given disease?", bold_lead="Incremental value: ")
BUL("because the same trainer runs on any cohort, does a score trained in one "
    "population transfer to another?", bold_lead="Transportability: ")
P("And at the atlas level, the causal engine answers: which immune proteins causally influence which "
  "diseases; in which direction (raising or lowering risk); whether the signal reflects a shared causal "
  "variant or LD; whether it holds at the protein as well as the transcript level; whether it replicates "
  "out of sample; which proteins act across multiple disease systems; and which of all these are novel, "
  "druggable and not already the target of an approved therapy.")

H2("I.3  How a research programme benefits")
BUL("the causal atlas provides a pre-filtered, "
    "direction-annotated, druggability-scored shortlist of immune targets, each tagged with the evidence "
    "tier it has earned, so scarce experimental resources are spent on the most credible hypotheses.",
    bold_lead="Target discovery: ")
BUL("the direction map states, per target, whether to block or to "
    "agonise, and the eQTL\u2013pQTL concordance flags whether a transcript signal survives at the protein "
    "level \u2014 the two facts most needed before committing to a modality.", bold_lead="Target validation: ")
BUL("recovering approved-drug axes with the correct direction is a "
    "genetic argument for repurposing and a check on the pipeline; the novelty engine then points beyond "
    "them.", bold_lead="Drug repurposing: ")
BUL("the PIRS can enrich trials for high-risk individuals, or "
    "identify a mechanistic subgroup, using a cheap plasma measurement.", bold_lead="Patient stratification: ")
BUL("because everything is built from public data with released code, "
    "any group can reproduce, extend to new diseases, or re-train the PIRS on their own cohort \u2014 the "
    "resource is a platform, not a static paper.", bold_lead="Reproducible foundation: ")

# ================= PART J: LIMITS / REPRO =================
H1("PART J \u00b7 Evidence tiers, limitations, reproducibility")

H2("J.1  The evidence ladder")
P("Every finding is bound to a tier, and language is bound to the tier. T2: nomination held for MHC/LD "
  "caution. T3: genetically-supported nomination (cis-MR FDR < 0.05). T4: prioritised transcript-level "
  "causal target (adds transcript colocalization PP.H4 \u2265 0.8). T5: protein-level causal target (adds a "
  "concordant protein-level pQTL-MR and protein colocalization). In the final integrated table this "
  "yields 2 T5, 15 T4, 8 T3 and 7 T2 findings. No result is described more strongly than its tier "
  "permits.")

H2("J.2  Limitations")
BUL("effects are estimated in blood, largely of European ancestry, from Finnish "
    "discovery GWAS. Replication is out-of-sample but still mostly European; ancestry generalisation is "
    "future work.", bold_lead="Population and tissue: ")
BUL("plasma protein validation rests on a single panel (INTERVAL/SomaScan); "
    "intracellular and MHC-region proteins have no plasma aptamer and cannot be validated this way.",
    bold_lead="Proteomic coverage: ")
BUL("single strongest-cis instruments cannot fully exclude horizontal "
    "pleiotropy; colocalization, protein-level triangulation and replication mitigate but do not "
    "eliminate it.", bold_lead="Instrument scope: ")
BUL("no nomination has yet been perturbed experimentally; the atlas ends at "
    "genetically-triangulated hypotheses, not proof of mechanism.", bold_lead="No functional perturbation: ")
BUL("the shipped resource is a causal atlas; the PIRS is a ready-to-run "
    "trainer, and its performance depends on the size, quality and event rate of the user\u2019s own cohort.",
    bold_lead="Predictive layer is user-trained: ")

H2("J.3  Reproducibility and data availability")
P("The pipeline is a set of numbered scripts run in order (proteome curation \u2192 instruments \u2192 cis-MR \u2192 "
  "colocalization \u2192 claim gate \u2192 protein-level pQTL \u2192 replication \u2192 supplementary figures \u2192 novelty "
  "\u2192 pan-phenome MR \u2192 pan-phenome analysis/coloc/novelty engine \u2192 figures \u2192 manuscript). Pre-computed "
  "result tables are released so figures and manuscripts regenerate without re-downloading multi-gigabyte "
  "summary statistics. All inputs are public (Olink Explore universe, Human Protein Atlas, MSigDB, "
  "eQTLGen, FinnGen R12, INTERVAL via the GWAS Catalog, OpenGWAS with a free token). Individual-level UK "
  "Biobank data and the UKB-PPP/deCODE gated pQTLs are controlled-access and were not used. Code is MIT; "
  "data tables and figures are CC-BY-4.0.")

# ================= PART K: PER-DISEASE RESULTS =================
H1("PART K \u00b7 Per-disease results: how the model predicts, and how it helps each disease")
P("This part turns the atlas from a global statistic into a disease-by-disease account. For every "
  "disease with at least one FDR-significant causal immune protein, we list the causal proteins (with "
  "direction of effect and whether the signal colocalizes to a shared causal variant), explain how those "
  "proteins drive prediction in a Plasma Immune Risk Score trained for that disease, and state concretely "
  "how the finding helps that disease \u2014 as a biomarker, a therapeutic hypothesis, or both. Notation: "
  "UP = genetically higher protein raises risk (a blockade/neutralisation target); DOWN = higher protein "
  "is protective (an agonism/replacement target); [coloc] = shared-variant colocalization PP.H4 \u2265 0.8. "
  "Across the 28-disease phenome the immune instruments produce 176 causal gene\u2013disease pairs; the "
  "leading proteins per disease are given below.")

def DIS(name, n, hits, predict, benefit):
    p = doc.add_paragraph(); r = p.add_run(f"{name}  "); r.bold = True; r.font.size = Pt(10.8)
    r2 = p.add_run(f"({n} causal immune protein{'s' if n!=1 else ''})"); r2.italic = True; r2.font.size = Pt(9.5); r2.font.color.rgb = GREY
    q = doc.add_paragraph(); rr = q.add_run("Causal immune proteins:  "); rr.bold = True; rr.font.size = Pt(9.5)
    q.add_run(hits).font.size = Pt(9.5)
    q2 = doc.add_paragraph(); rr = q2.add_run("Prediction:  "); rr.bold = True; rr.font.size = Pt(9.5)
    q2.add_run(predict).font.size = Pt(9.5)
    q3 = doc.add_paragraph(); rr = q3.add_run("How it helps:  "); rr.bold = True; rr.font.size = Pt(9.5)
    q3.add_run(benefit).font.size = Pt(9.5)
    q3.paragraph_format.space_after = Pt(9)

H2("K.1  Autoimmune diseases")
P("The autoimmune core is where the pipeline is best calibrated, because several hits are the targets of "
  "approved drugs and act as internal positive controls.")
DIS("Rheumatoid arthritis", 10,
    "IL6ST UP (OR 2.63) [coloc]; CTLA4 DOWN (OR 0.51) [coloc]; SWAP70 DOWN (OR 0.83) [coloc]; CA8 DOWN [coloc]; TNFRSF14 DOWN; C2 UP (MHC, held).",
    "A RA-PIRS is dominated by the IL-6 co-receptor axis (IL6ST) with positive weight and the "
    "checkpoint CTLA4 with protective weight; SWAP70 adds an independent, non-drug signal. Higher IL6ST "
    "and lower CTLA4 signalling both push predicted risk up.",
    "IL6ST recovers tocilizumab biology and CTLA4 recovers abatacept biology from genetics alone, "
    "validating the score; SWAP70 \u2014 the only novel RA target colocalizing at both transcript and protein "
    "and replicating out-of-sample \u2014 is a fresh, drug-naive hypothesis for RA.")
DIS("Type-1 diabetes (autoimmune endocrine)", 9,
    "IL2RA DOWN (OR 0.30) [coloc]; ERBB3 DOWN (OR 0.42) [coloc]; CTLA4 DOWN (OR 0.39) [coloc]; CTSH UP [coloc]; LIF UP; C2 UP (MHC).",
    "The score is driven by the IL-2 receptor-alpha (IL2RA/CD25) regulatory-T-cell axis and CTLA4, both "
    "protective, plus ERBB3; lower IL2RA/CTLA4 tone predicts higher risk.",
    "It reproduces the known IL2RA and CTLA4 autoimmune-diabetes loci (supporting low-dose IL-2 and "
    "checkpoint-modulating strategies) and nominates ERBB3 and CTSH as colocalized, testable T1D targets.")
DIS("Coeliac disease", 6,
    "HAVCR1/KIM-1 UP (OR 2.41) [coloc]; NUMB DOWN [coloc]; TGFA DOWN [coloc]; TNFSF8 DOWN; IL1RL1 DOWN; C2 UP (MHC).",
    "HAVCR1 carries strong positive weight; a coeliac-PIRS would flag individuals with high circulating "
    "KIM-1 as elevated-risk.",
    "HAVCR1/KIM-1 is a colocalized, drug-naive coeliac nomination and an accessible plasma biomarker; the "
    "protective TGFA and IL1RL1 (ST2) signals point to epithelial-repair and alarmin biology.")
DIS("Psoriasis", 6,
    "IL4 UP (OR 1.80); IFNGR2 DOWN (OR 0.88) [coloc]; IL6ST UP (OR 1.58) [coloc]; ERBB3 DOWN; SNX2 DOWN; C2 UP (MHC).",
    "IL4 and IL6ST push risk up; IFNGR2 (interferon-gamma receptor 2) is protective. The score blends a "
    "type-2 (IL4) and IL-6 signal.",
    "IL4 recovers dupilumab-relevant biology (direction-consistent), IL6ST supports IL-6-axis relevance, "
    "and IFNGR2 is a colocalized novel nomination \u2014 though its plasma-protein direction should be checked "
    "before it is read as a protein-level target.")
DIS("Multiple sclerosis", 2,
    "TNFSF14/LIGHT DOWN (OR 0.50) [coloc]; TNFRSF14/HVEM DOWN (OR 0.73).",
    "The MS-PIRS is anchored by the TNFSF14\u2013TNFRSF14 (LIGHT\u2013HVEM) co-stimulatory axis, both protective; "
    "lower LIGHT signalling predicts higher MS risk.",
    "TNFSF14\u2192MS is one of only two protein-level causal targets in the whole atlas (transcript + protein "
    "coloc, replicated at P = 2\u00d710\u207b\u00b9\u00b2 in the independent IMSGC GWAS) \u2014 the strongest single therapeutic "
    "hypothesis the resource produces.")
DIS("Ankylosing spondylitis", 5,
    "TNFRSF1A UP (OR 1.46) [coloc]; EDN1 UP (OR 2.02); RAB44 DOWN; KIAA0319 DOWN; C2 UP (MHC).",
    "TNFRSF1A (TNF receptor 1) drives positive weight; higher TNF-receptor signalling predicts higher risk.",
    "TNFRSF1A recovers anti-TNF (etanercept) biology from genetics with the correct blockade direction, "
    "validating the AS score and reinforcing the TNF axis as the central AS target.")
DIS("Sarcoidosis", 4,
    "CD226/DNAM-1 DOWN (OR 0.81) [coloc]; LTBR UP; ANXA11 DOWN; PVALB DOWN.",
    "CD226 (a T/NK activating receptor) is protective; the score weights immune-cell activation biology.",
    "CD226 is a colocalized novel sarcoidosis nomination linking NK/T-cell adhesion to granuloma biology; "
    "LTBR points to lymphotoxin-beta-receptor signalling.")
DIS("Sj\u00f6gren syndrome", 2,
    "PLAU/uPA UP (OR 1.66) [coloc]; C2 UP (MHC, held).",
    "PLAU (urokinase plasminogen activator) carries the predictive signal outside the MHC.",
    "PLAU is a colocalized novel Sj\u00f6gren nomination implicating plasminogen-activation/tissue-remodelling "
    "biology; note Sj\u00f6gren has no FinnGen-independent GWAS, so this remains a transcript-level nomination.")
DIS("Guillain\u2013Barr\u00e9 syndrome", 2,
    "SIGLEC7 UP (OR 3.05) [coloc]; SIGLEC9 UP (OR 2.94) [coloc].",
    "Both sialic-acid-binding inhibitory receptors carry strong positive weight; a GBS-PIRS would be "
    "dominated by the SIGLEC7/9 pair.",
    "SIGLEC7/9\u2192GBS is a striking, colocalized, entirely drug-naive nomination that connects "
    "glycan-recognition immune-inhibitory receptors to an acute demyelinating neuropathy \u2014 a novel axis "
    "worth functional follow-up (untestable in replication for lack of an independent GBS GWAS).")
DIS("Crohn\u2019s disease", 2,
    "HNMT DOWN (OR 0.64) [coloc]; BCL7A DOWN (OR 0.57) [coloc].",
    "Histamine N-methyltransferase (HNMT) is protective; lower HNMT predicts higher risk.",
    "HNMT is a colocalized novel Crohn\u2019s nomination implicating histamine catabolism in gut inflammation "
    "\u2014 a mechanistically fresh, testable hypothesis.")
DIS("Autoimmune hyperthyroidism / Systemic lupus / Vitiligo", 3,
    "CTLA4 DOWN [coloc] and IL12RB2 UP [coloc] (hyperthyroidism); C2 UP (lupus, MHC); (vitiligo signals MHC-dominated).",
    "CTLA4 anchors the thyroid-autoimmunity score; lupus is dominated by the held MHC C2 signal.",
    "CTLA4 again recovers checkpoint biology (validating), while the recurrent C2/MHC signal across five "
    "autoimmune diseases is explicitly flagged as LD-confounded and withheld from causal claims.")

H2("K.2  Cardiovascular diseases")
P("The pan-phenome expansion reveals that immune proteins act causally well beyond autoimmunity, most "
  "densely in cardiovascular disease (67 hits).")
DIS("Hypertension", 28,
    "ACE UP (OR 1.61); FES DOWN [coloc]; IFNGR2 UP (OR 1.07) [coloc]; MERTK UP [coloc]; SWAP70 DOWN; SPINK8 DOWN.",
    "The largest causal set in the phenome: the hypertension-PIRS integrates ACE (positive), the "
    "granulocyte kinase FES (protective), and the interferon and efferocytosis receptors IFNGR2/MERTK.",
    "ACE recovers a validated antihypertensive target (ACE inhibitors), confirming the score; IFNGR2 and "
    "MERTK are colocalized novel immune\u2013cardiovascular nominations connecting interferon and macrophage "
    "clearance biology to blood-pressure regulation.")
DIS("Coronary heart disease", 15,
    "PSRC1 DOWN (OR 0.85) [coloc]; PLAUR UP (OR 1.59) [coloc]; SCARB1 DOWN [coloc]; LPL DOWN; FES DOWN.",
    "A CHD-PIRS weights the PSRC1/SORT1-region lipid axis and the urokinase receptor PLAUR; PLAUR raises "
    "predicted risk, PSRC1 and the HDL receptor SCARB1 lower it.",
    "PSRC1 and SCARB1 recover known lipid/CHD biology (calibrating the score), while PLAUR is a "
    "colocalized novel inflammatory-thrombosis CHD nomination \u2014 a druggable receptor linking innate "
    "immunity to atherothrombosis.")
DIS("Atrial fibrillation", 12,
    "ACE UP (OR 1.72) [coloc]; FES DOWN [coloc]; PFKFB2 DOWN [coloc]; CDKN1A UP; ADAM15 DOWN.",
    "ACE and the cell-cycle inhibitor CDKN1A drive positive weight; FES and the glycolytic PFKFB2 are "
    "protective.",
    "ACE colocalizes with AF as well as hypertension, suggesting a shared renin\u2013angiotensin axis; PFKFB2 "
    "and CDKN1A are novel colocalized nominations linking metabolic/senescence biology to arrhythmia.")
DIS("Venous thromboembolism", 8,
    "PROCR DOWN (OR 0.55); PDLIM7 DOWN [coloc]; SCPEP1 UP [coloc]; TSPAN15 DOWN; TGFB2 UP.",
    "The VTE-PIRS is anchored by the endothelial protein-C receptor PROCR (strongly protective); lower "
    "PROCR predicts higher clot risk.",
    "PROCR recovers established coagulation biology (protein-C pathway), validating the score; PDLIM7 and "
    "SCPEP1 are colocalized novel VTE nominations.")
DIS("Heart failure", 4,
    "CDKN1A UP (OR 1.48); PFKFB2 DOWN; RAB44 DOWN; GSTM4 UP.",
    "The senescence marker CDKN1A/p21 drives positive weight; higher CDKN1A predicts higher heart-failure "
    "risk.",
    "CDKN1A links cellular senescence to heart failure \u2014 a novel, aging-relevant immune\u2013cardiac "
    "hypothesis that dovetails with the neuro/aging category.")

H2("K.3  Metabolic diseases")
DIS("Type-2 diabetes", 19,
    "SPINK8 DOWN (OR 0.83) [coloc]; ARG1 DOWN; ANPEP UP; SPRY2 UP; NOMO1 UP.",
    "A T2D-PIRS integrates the protease-inhibitor SPINK8 (protective) with aminopeptidase and "
    "sprouty-signalling proteins; the immune feature set adds inflammatory context to metabolic risk.",
    "SPINK8 is a colocalized novel T2D nomination and, because it also acts on hypertension, a candidate "
    "shared cardiometabolic control point.")
DIS("Obesity", 7,
    "F12 UP (OR 1.48) [coloc]; LTBP3 UP [coloc]; SERPING1 UP; GDF15 DOWN; MAPRE3 DOWN.",
    "Coagulation factor XII (F12) and the TGF-beta binding protein LTBP3 drive positive weight; GDF15 is "
    "protective.",
    "F12 and LTBP3 are colocalized novel obesity nominations linking contact-activation and TGF-beta "
    "biology to adiposity; GDF15 recovers a known appetite/energy-balance axis.")
DIS("Type-1 diabetes (metabolic view)", 9,
    "(see Autoimmune) IL2RA/CTLA4/ERBB3 DOWN [coloc]; CTSH UP [coloc].",
    "Same regulatory-T-cell and ERBB3 axes as above, here scored against the metabolic-category label.",
    "Reinforces that T1D risk is immune-driven and that IL2RA/CTLA4 are the most credible modulatable "
    "axes.")

H2("K.4  Neurodegenerative and aging diseases")
DIS("Dementia (all-cause)", 14,
    "ACE DOWN (OR 0.44) [coloc]; RAPGEF2 UP [coloc]; LYPD3 DOWN; PVR DOWN; RELB DOWN; TEX101 UP.",
    "A dementia-PIRS is anchored by ACE (here protective, opposite to its hypertension direction), with "
    "PVR (the poliovirus receptor/CD155) and LYPD3 contributing.",
    "ACE colocalizes with dementia AND Alzheimer\u2019s at PP.H4 > 0.97 \u2014 a top novel-priority target linking "
    "the renin\u2013angiotensin system to neurodegeneration, with a direction opposite to its cardiovascular "
    "effect (a caution for any systemic ACE-modulating strategy).")
DIS("Alzheimer\u2019s disease", 9,
    "LYPD3 DOWN (OR 0.24); ACE DOWN (OR 0.43) [coloc]; RAPGEF2 UP [coloc]; PVR DOWN; TEX101 UP; CRIM1 UP.",
    "LYPD3 carries the strongest single effect; ACE and RAPGEF2 colocalize. The AD-PIRS is largely shared "
    "with the dementia score.",
    "The shared ACE and LYPD3/PVR signals nominate immune-adjacent neurodegeneration axes; PVR/CD155 "
    "connects an immune checkpoint ligand to AD, a novel and druggable direction.")
DIS("Glaucoma", 5,
    "LTBP2 DOWN (OR 0.67); SPON2 UP; CRISP3 UP; ANXA11 UP; TNFSF13 DOWN.",
    "The extracellular-matrix protein LTBP2 (protective) anchors the glaucoma score.",
    "LTBP2 recovers known ocular extracellular-matrix biology; TNFSF13 (APRIL) adds a testable immune "
    "nomination.")
DIS("Epilepsy / Osteoporosis", 3,
    "PM20D1 DOWN (OR 0.87) [coloc] (epilepsy); SEZ6 DOWN, DPEP2 DOWN (osteoporosis).",
    "PM20D1 anchors the epilepsy signal and, spanning three disease categories, is a highly pleiotropic "
    "control point.",
    "PM20D1 is a colocalized novel nomination acting across epilepsy, obesity and type-1 diabetes \u2014 a "
    "genuinely multi-system immune-metabolic axis surfaced by the pan-phenome view.")

H2("K.5  Renal disease")
DIS("Chronic kidney disease", 1,
    "AAMDC DOWN (OR 0.88).",
    "A CKD immune signal is sparse; AAMDC provides the single causal nomination.",
    "AAMDC (which also appears in venous thromboembolism) is a preliminary CKD nomination; the sparse "
    "renal signal is reported honestly rather than inflated.")

# ================= PART L: EXTENDED NOVELTY =================
H1("PART L \u00b7 Extended novelty: dossiers, therapeutic modality and mechanistic hypotheses")
P("This part consolidates what is genuinely new in the resource. Novelty here does not mean an "
  "unprecedented statistic; it means a target, direction or axis that (i) is supported by convergent "
  "genetic evidence, (ii) is not already the object of an approved drug for that indication, and (iii) "
  "carries an actionable, direction-specified therapeutic implication. The novelty-priority engine "
  "(Part G.3) formalises this by adding causal, colocalization, pleiotropy, druggability and cell-source "
  "evidence and subtracting known-drug and MHC penalties, leaving 45 novel colocalized targets. Below we "
  "expand the leading nominations into dossiers, map them to therapeutic modality, and articulate the "
  "cross-system hypotheses that only a pan-phenome causal atlas can generate.")

H2("L.1  Novel-target dossiers")
P("Each dossier states the target, indication, direction (and therefore modality), the convergent "
  "evidence, the cell source, and the single most important caveat.")
def DOSS(target, body):
    p = doc.add_paragraph(); r = p.add_run(target); r.bold = True; r.font.size = Pt(10.3); r.font.color.rgb = BLUE
    q = doc.add_paragraph(); q.add_run(body).font.size = Pt(9.5); q.paragraph_format.space_after = Pt(8)
DOSS("SWAP70 \u2014 rheumatoid arthritis (agonise; higher protein protective)",
     "The single strongest novel nomination: the only new target colocalizing at BOTH transcript and "
     "protein level (PP.H4 > 0.97) AND replicating in the independent Okada GWAS. SWAP70 is a "
     "B-cell/actin-regulatory protein; the genetics argue that raising or mimicking its function is "
     "protective in RA. Caveat: mechanism of the protective direction needs functional confirmation.")
DOSS("SIGLEC7 & SIGLEC9 \u2014 Guillain\u2013Barr\u00e9 syndrome (block; higher protein raises risk)",
     "A paired, colocalized nomination (OR ~3) implicating sialic-acid-binding inhibitory receptors in "
     "acute demyelinating neuropathy \u2014 a glycoimmunology axis with no approved GBS drug. Druggable as "
     "cell-surface receptors. Caveat: no independent GBS GWAS exists, so this stays transcript-level.")
DOSS("ACE \u2014 dementia / Alzheimer\u2019s disease (context-dependent)",
     "ACE colocalizes with both dementia and Alzheimer\u2019s (PP.H4 > 0.97) with a PROTECTIVE direction, "
     "opposite to its risk-raising effect in hypertension/atrial fibrillation. This directional split is "
     "itself the novelty and the caution: a systemic ACE inhibitor beneficial for blood pressure could be "
     "directionally unfavourable for neurodegeneration. A prime example of why direction, not just "
     "identity, matters.")
DOSS("IFNGR2 \u2014 hypertension & psoriasis (cross-category)",
     "Colocalized in both a cardiovascular and an autoimmune disease, with opposite directions "
     "(risk-raising in hypertension, protective in psoriasis), marking the interferon-gamma receptor as a "
     "pleiotropic immune\u2013cardiovascular control point. Caveat: plasma-protein direction should be "
     "confirmed before modality is fixed.")
DOSS("PLAUR \u2014 coronary heart disease (block; higher protein raises risk)",
     "The urokinase receptor colocalizes with CHD (OR 1.59), a druggable cell-surface receptor linking "
     "innate-immune proteolysis to atherothrombosis \u2014 a novel, drug-naive CHD target.")
DOSS("ERBB3 \u2014 type-1 diabetes (agonise; higher protein protective)",
     "Colocalized T1D nomination (PP.H4 0.90) independent of the classical checkpoint loci, implicating "
     "receptor-tyrosine-kinase biology in beta-cell autoimmunity.")
DOSS("SPINK8 & PM20D1 \u2014 shared cardiometabolic / multi-system control points",
     "SPINK8 colocalizes with type-2 diabetes and acts on hypertension; PM20D1 acts across epilepsy, "
     "obesity and type-1 diabetes (three categories). Both are novel, druggability-flagged, and exemplify "
     "targets that a single-disease study could never have prioritised.")
DOSS("HAVCR1 / KIM-1 \u2014 coeliac disease (block; higher protein raises risk)",
     "A colocalized coeliac nomination that is also an accessible plasma biomarker (KIM-1), giving a rare "
     "target that is simultaneously a candidate drug target and a monitoring assay.")

H2("L.2  Direction-aware therapeutic modality map")
P("Because every causal estimate carries a sign, each target maps to a modality without further "
  "assumption. This is a central novelty of the resource: it does not merely list targets, it prescribes "
  "whether to raise or lower each one.")
TABLE(["Direction (genetics)","Interpretation","Therapeutic modality","Example novel targets"],
      [["OR > 1 (protein raises risk)","protein is causal & harmful","block / neutralise / antagonise (antibody, small molecule)","PLAUR (CHD), HAVCR1 (coeliac), SIGLEC7/9 (GBS), IFNGR2 (hypertension)"],
       ["OR < 1 (protein is protective)","protein is causal & beneficial","agonise / replace / stabilise (recombinant, agonist)","SWAP70 (RA), ERBB3 (T1D), PROCR (VTE), SPINK8 (T2D)"],
       ["opposite across diseases","pleiotropic, context-dependent","indication-specific; systemic modulation risky","ACE (CV vs dementia), IFNGR2 (hypertension vs psoriasis)"]])
P("The druggability score refines this: secreted proteins favour recombinant/agonist or "
  "neutralising-antibody strategies, single-pass membrane receptors favour antagonist antibodies, and "
  "FDA-target-class flags mark targets with established tractability. Combining direction with "
  "druggability yields a per-target modality recommendation rather than a bare hit list.")

H2("L.3  Cross-system immune-axis hypotheses")
P("The pan-phenome design surfaces hypotheses invisible to any single-disease analysis \u2014 immune "
  "proteins that act as control points across organ systems:")
BUL("the interferon-gamma receptor links autoimmune (psoriasis) and cardiovascular "
    "(hypertension) disease with opposite signs, suggesting interferon tone is a shared, "
    "context-dependent axis.", bold_lead="IFNGR2 (interferon \u2194 vascular): ")
BUL("the renin\u2013angiotensin enzyme is causally protective in "
    "neurodegeneration yet risk-raising in hypertension and atrial fibrillation \u2014 a systemic "
    "trade-off with direct implications for lifelong ACE-inhibitor use.", bold_lead="ACE (vascular \u2194 brain): ")
BUL("the Mer/efferocytosis receptor and myeloperoxidase connect "
    "macrophage-clearance and oxidative biology across cardiovascular and metabolic disease.",
    bold_lead="MERTK & MPO (immune clearance \u2194 cardiometabolic): ")
BUL("a single metabolic-immune protein acting across epilepsy, obesity and "
    "type-1 diabetes \u2014 the most pleiotropic control point in the atlas.", bold_lead="PM20D1 (three categories): ")
BUL("a B-cell/actin regulator protective in both rheumatoid arthritis and "
    "hypertension, colocalized in RA \u2014 a shared autoimmune\u2013vascular node.", bold_lead="SWAP70 (autoimmune \u2194 vascular): ")

H2("L.4  What is novel relative to prior work")
P("Prior plasma-proteomic disease studies are typically (i) association-based rather than causal, (ii) "
  "confined to a single disease or disease family, (iii) reliant on gated individual-level cohorts, and "
  "(iv) silent on therapeutic direction. This resource is novel on all four axes simultaneously: it is "
  "genetically causal (MR + colocalization + protein-level validation + independent replication), it "
  "spans 28 diseases in five categories, it is built entirely from public data and is fully "
  "reproducible, and it assigns every target an explicit therapeutic direction and a druggability-aware "
  "modality. The integrated novelty-priority engine then does something no prior plasma-immune study "
  "does: it ranks targets by cumulative orthogonal evidence while actively penalising the approved-drug "
  "and MHC axes, so the output is a prioritised list of genuinely new, direction-specified, druggable "
  "immune targets rather than a re-statement of known biology.")

# ================= PART M: INTELLIGENCE LAYER =================
H1("PART M \u00b7 The disease-trained PIRS intelligence layer")
P("The final component of the resource is an intelligence layer that sits on top of both the trained "
  "PIRS and the genetic causal atlas and turns their raw outputs into a single, ranked plasma-immune "
  "discovery report (src/31). Its purpose is to answer the question a translational reader actually asks: "
  "given a disease and its plasma immune proteome, which proteins are worth acting on, in which "
  "direction, with what confidence, and what is the next experiment? It does not return a bare risk "
  "score or an unannotated protein list.")

H2("M.1  Inputs and mode")
P("The layer reads six real tables \u2014 the curated plasma-immune annotation, the pan-phenome cis-MR "
  "results, the pan-phenome colocalization posteriors, the integrated novelty-engine ranking, the final "
  "evidence-tier table with protein-level pQTL and replication status, and the protein-level pQTL MR "
  "results \u2014 plus any trained PIRS weight and cross-validation files. When a PIRS has been trained on an "
  "authorised cohort, its coefficients, cross-validated stability and discrimination contributions are "
  "fused in. When no PIRS is present (the default distributed state, because no individual-level data "
  "ships with the resource), the layer runs in causal-atlas-only mode: every causal, novelty, direction "
  "and tier field is populated from real genetics, and the predictive fields are written as "
  "\u201cNA (train PIRS)\u201d rather than invented \u2014 the same no-fabrication discipline that governs the whole "
  "project.")

H2("M.2  Causal\u2013predictive concordance")
P("For each gene\u2013disease pair the layer classifies the agreement between the predictive signal (the sign "
  "of the PIRS weight) and the causal signal (the MR odds ratio). A PIRS weight and an MR effect that "
  "point the same way, backed by colocalization or a concordant plasma pQTL, is the strongest form of "
  "internal corroboration: the protein both predicts and appears to cause the disease. In causal-only "
  "mode the field records the causal direction and marks the predictive comparison as pending a trained "
  "model.")

H2("M.3  The Plasma Immune Novelty Score and tiers")
P("The layer computes a bounded Plasma Immune Novelty Score (PINS) that rewards causal strength "
  "(\u2212log10 FDR), colocalization posterior, a protein-level pQTL bonus, druggability and cross-disease "
  "pleiotropy, while penalising the approved-drug and MHC axes, and then assigns a novelty tier:")
BUL("known-drug positive control \u2014 recovered from genetics, validating the pipeline rather than nominating "
    "a new target.", bold_lead="Tier 1  ")
BUL("an MHC/LD-flagged signal held at nomination until classical HLA alleles are excluded.", bold_lead="Tier 2  ")
BUL("a causal nomination from cis-MR that has not yet colocalized.", bold_lead="Tier 3  ")
BUL("a prioritised causal target \u2014 cis-MR plus colocalization \u2014 that is novel.", bold_lead="Tier 4  ")
BUL("a novel plasma-immune target that reaches the full bar: prediction-ready, colocalized, protein-level "
    "pQTL-concordant, specific and druggable.", bold_lead="Tier 5  ")
P("Only tiers 4 and 5 count as high-novelty. Applied to the 176 causal gene\u2013disease pairs across 25 "
  "diseases, the layer returns 45 Tier-4 prioritised targets, 12 Tier-1 positive controls and 7 "
  "MHC-held signals, with 6 pairs reaching protein-level causality (transcript colocalization plus a "
  "direction-concordant INTERVAL plasma pQTL) and 17 independently replicated.")

H2("M.4  Why Tier 5 is currently empty \u2014 an honest gate")
P("No pair currently occupies Tier 5, and this is a truthful data-coverage boundary rather than a gap in "
  "the run. Tier 5 demands all five criteria at once. The two protein-level causal hits, TNFSF14 in "
  "multiple sclerosis and SWAP70 in rheumatoid arthritis, each miss one: TNFSF14 is flagged as a known-"
  "drug axis and is therefore reported as a Tier-1 positive control, while SWAP70 is genuinely novel and "
  "protein-level causal but has druggability zero (an intracellular, non-secreted protein) and so lands "
  "at Tier 4. Separately, the cardiovascular, metabolic, renal and neurological arms of the phenome have "
  "no plasma pQTL layer computed yet \u2014 INTERVAL pQTL was colocalized only against the autoimmune arc \u2014 "
  "so none of those hits can reach a protein-level tier until a plasma pQTL panel is colocalized against "
  "them. Colocalizing a plasma pQTL resource across the full phenome is the single step that would "
  "promote Tier-4 targets to Tier 5; the layer names this explicitly rather than manufacturing a "
  "Tier-5 ceiling.")

H2("M.5  Therapeutic direction and the Final Required Output Table")
P("Every causal target is assigned a therapeutic modality from its effect direction and localisation: a "
  "risk-raising protein (OR>1) is a blockade target (antagonist, neutralising antibody or small "
  "molecule); a protective protein (OR<1) is an agonism or replacement target, with soluble receptors "
  "routed to decoy/replacement strategies and secreted proteins to recombinant supplementation. Pairs "
  "that do not reach a causal tier are labelled biomarker-only. All of this is emitted as the Final "
  "Required Output Table (T6), a ranked, roughly thirty-column sheet carrying, for each target, its "
  "disease, protein and gene, protein class and plasma detectability, PIRS coefficient and direction "
  "(or the train-PIRS placeholder), cross-validated stability and sensitivity/specificity/AUROC "
  "contributions, the full MR/coloc/pQTL/replication evidence, known-drug status and druggability, the "
  "novelty score and tier, the causal\u2013predictive concordance call, the therapeutic direction, a "
  "biomarker-versus-target classification, and \u2014 crucially for a reader planning work \u2014 the best figure "
  "panel, the best next validation experiment and a final recommendation.")

H2("M.6  Six figure panels and claim discipline")
P("The layer renders six panels: a workflow schematic (A), disease-specific PIRS performance or, "
  "untrained, the per-disease causal-signal strength (B), the direction-coloured plasma-immune signature "
  "of the high-novelty targets (C), the causal-evidence concordance ladder from MR to colocalization to "
  "protein-level pQTL to replication (D), the novelty-priority map of PP.H4 against PINS (E), and a "
  "validation plan for the leading novel targets (F). Throughout, claims are bound to the evidence "
  "actually reached: a PIRS weight alone is a biomarker; adding cis-MR makes it a causal nomination; "
  "adding colocalization makes it a prioritised causal target; adding a direction-concordant plasma "
  "pQTL makes it a protein-level causal target; and only experimental perturbation would make it a proof "
  "of mechanism \u2014 which the layer proposes as the next experiment rather than asserts. All statements "
  "remain restricted to plasma immune proteins; no single-cell, tissue or intracellular claim is made "
  "unless such data are separately supplied.")

# ================= PART N: PROOF OF CONCEPT =================
H1("PART N \u00b7 Proof of concept: does the model work?")
P("A genetics-anchored target-discovery pipeline is credible only if it independently "
  "rediscovers biology that is already validated in humans \u2014 approved drug targets \u2014 from "
  "genetics alone, in the correct pharmacological direction, with no drug information supplied. "
  "This is the honest positive-control test, and the atlas passes it.")

H2("N.1  Approved-drug axes recovered with the correct direction")
P("Five gene\u2013disease pairs that correspond to marketed immune drugs are recovered by the "
  "discovery engine, each in the direction the drug actually acts:")
TABLE(
    ["Target \u2192 disease", "Model effect (OR)", "Model action", "Coloc PP.H4",
     "Replication P", "Approved drug (mechanism)"],
    [["IL6ST \u2192 rheumatoid arthritis", "2.63", "block", "1.00", "7\u00d710\u207b\u00b2\u2074",
      "tocilizumab (IL-6R blockade)"],
     ["CTLA4 \u2192 rheumatoid arthritis", "0.51", "agonize", "0.98", "3\u00d710\u207b\u00b2\u2070",
      "abatacept (CTLA4 co-stim agonist)"],
     ["CTLA4 \u2192 autoimmune hyperthyroidism", "0.15", "agonize", "0.84", "4\u00d710\u207b\u00b9\u2079",
      "abatacept (CTLA4 co-stim agonist)"],
     ["TNFRSF1A \u2192 ankylosing spondylitis", "1.46", "block", "0.86", "6\u00d710\u207b\u2076",
      "etanercept (TNF blockade)"],
     ["IL4 \u2192 psoriasis", "1.80", "block", "\u2014", "4\u00d710\u207b\u2074",
      "dupilumab (IL-4R\u03b1 blockade)"]])
P("A noise pipeline would not preferentially rank the exact proteins pharma has already "
  "validated in humans, nor infer whether to block or agonize each one. That it does both, and "
  "that the same signals colocalize and replicate in independent non-FinnGen GWAS, is the "
  "evidence that the discovery engine is calibrated.")

H2("N.2  Cardiovascular spotlight \u2014 ACE")
P("The cleanest cardiovascular positive control is ACE, the target of ACE-inhibitors (ramipril, "
  "lisinopril) \u2014 the most-prescribed cardiovascular drug class. From genetics alone the model "
  "infers ACE \u2192 hypertension OR \u2248 1.61 (FDR 7\u00d710\u207b\u2079) and ACE \u2192 atrial fibrillation OR \u2248 1.72, "
  "both risk-raising \u2192 block, and colocalizes the signal (PP.H4 up to 0.88). ACE-inhibitors work "
  "by blocking ACE to lower blood pressure: the model reaches the same target and the same "
  "direction, with no cardiovascular-drug knowledge supplied.")

IMG("08_figures/intelligence_layer/PROOF_validation_exhibit.png", width_in=6.7,
    caption="Validation exhibit. (A) Five approved-drug axes recovered from genetics alone in "
            "the correct pharmacological direction (red = risk\u2192block, green = protective\u2192"
            "agonize). (B) Cardiovascular spotlight: ACE causal odds ratios for hypertension and "
            "atrial fibrillation \u2014 the ACE-inhibitor target and direction. (C) Orthogonal support: "
            "colocalization PP.H4 and independent-GWAS replication \u2212log\u2081\u2080P for each control.")

H2("N.3  What this proves \u2014 and what it does not")
P("It proves the discovery engine (cis-MR \u2192 colocalization \u2192 replication) is calibrated across "
  "both autoimmune and cardiovascular disease, so the novel targets it ranks sit on the same "
  "evidence scale as these controls. It does not yet prove protein-level causality for the "
  "cardiovascular hits: ACE and the other pan-phenome signals are transcript-level (tier 4) \u2014 "
  "colocalized but without a plasma pQTL layer. ACE\u2019s drug validation is external proof; making "
  "ACE a tier-5 protein-level causal target inside the pipeline would require colocalizing a "
  "plasma ACE pQTL against these diseases. No claim is made beyond the evidence actually reached.")

# ================= WORKED EXAMPLE =================
H1("Appendix 1 \u00b7 A worked interpretation, end to end")
P("To make the pipeline concrete, follow one target \u2014 SWAP70 in rheumatoid arthritis \u2014 through every "
  "layer, exactly as the resource does.")
NUM("SWAP70 passes the plasma-immunome inclusion score (a B-cell-lineage immune protein) and enters the "
    "1,007-protein panel.", bold_lead="Curation. ")
NUM("its strongest blood cis-eQTL in eQTLGen is selected; the Z-score is converted to a standardised "
    "beta and standard error by the Zhu transformation.", bold_lead="Instrument. ")
NUM("harmonised against the FinnGen rheumatoid-arthritis GWAS, the Wald ratio gives OR \u2248 0.83 \u2014 "
    "genetically higher SWAP70 is protective \u2014 surviving 5% FDR.", bold_lead="cis-MR. ")
NUM("coloc.abf returns PP.H4 > 0.96, rejecting the two-distinct-variant hypothesis: the protein and "
    "disease signals share one causal variant. SWAP70 is promoted to a transcript-level causal target "
    "(tier 4).", bold_lead="Colocalization. ")
NUM("an INTERVAL plasma cis-pQTL for SWAP70 gives a directionally concordant protein-level MR and "
    "protein colocalization PP.H4 > 0.97, promoting SWAP70 to a protein-level causal target (tier 5).",
    bold_lead="Protein validation. ")
NUM("the instrument variant is directionally concordant in the independent Okada rheumatoid-arthritis "
    "GWAS, so the nomination is not a Finnish artefact.", bold_lead="Replication. ")
NUM("SWAP70 carries no approved-drug penalty and no MHC penalty, colocalizes, and is druggability-"
    "flagged, so the engine ranks it as the leading NOVEL rheumatoid-arthritis target; its protective "
    "direction specifies an agonism/replacement modality.", bold_lead="Novelty engine. ")
NUM("in a rheumatoid-arthritis PIRS, SWAP70 contributes a protective (negative-risk) weight; lower "
    "plasma SWAP70 raises predicted risk, and the weight is interpretable because the same protein is a "
    "tier-5 causal target.", bold_lead="Prediction. ")
P("The same eight steps, applied across 812 genes and 28 diseases, generate the entire atlas; the "
  "evidence tier a target reaches is simply how far along this chain its data carry it.")

# ================= FORMULAE =================
H1("Appendix 2 \u00b7 Statistical formulae summary")
P("Effect-size reconstruction from an eQTL Z-score (Zhu et al., 2016), with effect-allele frequency p "
  "and sample size n:")
EQ("denom = 2 p (1\u2212p) (n + Z\u00b2);   beta = Z / \u221adenom;   se = 1 / \u221adenom")
P("Single-instrument causal effect (Wald ratio) and its standard error:")
EQ("beta_MR = beta_out / beta_exp;   se_MR = se_out / |beta_exp|;   OR = exp(beta_MR)")
P("Benjamini\u2013Hochberg FDR: order p-values p(1)\u2264\u2026\u2264p(m); reject where p(i) \u2264 (i/m)\u00b7q, q = 0.05.")
P("Per-SNP Wakefield approximate Bayes factor (log space), variance V = se\u00b2, prior variance W:")
EQ("r = W / (W+V);   lABF = 0.5 ( log(1\u2212r) + r Z\u00b2 )")
P("Colocalization posteriors combine per-hypothesis summed Bayes factors with priors p1, p2, p12; "
  "PP.H4 is the posterior of one shared causal variant. Promotion threshold PP.H4 \u2265 0.8.")
P("Novelty-priority score (per gene\u2013disease pair):")
EQ("NP = min(\u2212log10FDR/5, 2) + PP.H4 + 0.5(n_cat\u22121) + s_drug + 0.5 \u2212 1.5\u00b7known \u2212 2.0\u00b7MHC")
P("Elastic-net Cox objective (partial-likelihood loss L with mixing \u03b1 = l1_ratio, penalty \u03bb):")
EQ("min \u2212L(\u03b2) + \u03bb [ \u03b1 \u2016\u03b2\u20161 + \u00bd(1\u2212\u03b1) \u2016\u03b2\u2016\u00b2\u2082 ]")
P("Discrimination: Harrell\u2019s C-index (survival) or AUROC (fixed-horizon logistic), estimated by "
  "stratified k-fold cross-validation with fold-internal preprocessing.")

# ================= GLOSSARY =================
H1("Appendix 3 \u00b7 Glossary")
gloss = [
 ("NPX","Normalised Protein eXpression \u2014 Olink\u2019s relative log2 protein-abundance unit."),
 ("cis-eQTL","a variant near a gene that is associated with that gene\u2019s expression."),
 ("cis-pQTL","a variant near a gene that is associated with its plasma protein level."),
 ("MR","Mendelian randomization \u2014 causal inference using genetic variants as instruments."),
 ("Wald ratio","single-instrument MR estimate = outcome effect / exposure effect."),
 ("OR","odds ratio; OR>1 = higher protein raises risk, OR<1 = lowers risk."),
 ("FDR","false-discovery rate; controlled here at 5% by Benjamini\u2013Hochberg."),
 ("coloc / PP.H4","colocalization; posterior probability that exposure and outcome share one causal variant."),
 ("Wakefield ABF","per-SNP approximate Bayes factor used inside coloc.abf."),
 ("LD","linkage disequilibrium \u2014 correlation between nearby variants; extreme in the MHC."),
 ("MHC","major histocompatibility complex (chromosome 6); LD-confounded, held at nomination."),
 ("elastic net","penalised regression mixing L1 (selection) and L2 (stability)."),
 ("C-index","Harrell\u2019s concordance for right-censored survival discrimination."),
 ("AUROC","area under the ROC curve; fixed-horizon classification discrimination."),
 ("PIRS","Plasma Immune Risk Score \u2014 the supervised model of this document."),
 ("evidence tier","T2 MHC-caution \u2192 T3 nomination \u2192 T4 transcript-causal \u2192 T5 protein-causal."),
]
TABLE(["Term","Definition"], gloss)

# ================= REFERENCES =================
H1("Appendix 4 \u00b7 Key references and resources")
for ref in [
 "Sun BB et al. Genomic atlas of the human plasma proteome (INTERVAL). Nature, 2018.",
 "V\u00f5sa U et al. Large-scale cis- and trans-eQTL analyses (eQTLGen). Nat Genet, 2021.",
 "Kurki MI et al. FinnGen: genome-wide association of the Finnish biobank (R12). Nature, 2023.",
 "Giambartolomei C et al. Bayesian colocalization (coloc.abf). PLoS Genet, 2014.",
 "Wakefield J. Bayes factors for genome-wide association studies. Genet Epidemiol, 2009.",
 "Zhu Z et al. Integration of summary data (SMR); Z-to-effect transformation. Nat Genet, 2016.",
 "Uhl\u00e9n M et al. Human Protein Atlas: tissue and immune-cell proteome. Science, 2015/2019.",
 "Liberzon A et al. Molecular Signatures Database (MSigDB C7/C8). Cell Systems, 2015.",
 "Hemani G et al. The MR-Base / OpenGWAS platform. eLife, 2018.",
 "Benjamini Y, Hochberg Y. Controlling the false discovery rate. JRSS-B, 1995.",
 "Simon N et al. Regularization paths for Cox\u2019s model via coordinate descent (elastic-net Cox). J Stat Softw, 2011.",
]:
    p = doc.add_paragraph(style="List Number"); p.add_run(ref).font.size = Pt(9.5)

# closing
doc.add_paragraph()
c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run("End of methodology monograph."); r.italic = True; r.font.color.rgb = GREY

out = os.path.join(OUT, "PIRS_and_Atlas_Methodology.docx")
doc.save(out)
print("wrote", out)

# quick page/paragraph diagnostics
import zipfile, re
data = open(out,"rb").read()
print("size MB:", round(len(data)/1e6,2))
z = zipfile.ZipFile(out); xml = z.read("word/document.xml").decode("utf8","ignore")
print("paragraphs:", xml.count("<w:p ")+xml.count("<w:p>"), "| headings:", xml.count("Heading"))
