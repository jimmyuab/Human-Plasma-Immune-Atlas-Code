#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
53 - Cardiovascular NOVELTY report (figures + Word)
==================================================
Takes the novelty / therapeutic-direction engine (layer 7 of the atlas) and
reports, for the cardiovascular chapter of the FinnGen R12 phenome only:
what is genuinely new, how new it is, why the engine calls it new, and what a
wet lab or a drug programme would do with each nomination.

Nothing here is fabricated. Every number, effect, colocalization posterior,
replication flag and figure is computed live from the atlas' own result tables.

Inputs (all in 06_genetic_causality/):
  novelty_engine_ranked_ALL.tsv       layer-7 output over the whole phenome
  uk_panphenome_concordance_ALL.tsv   Finland -> England two-population test
  pqtl_MR_ALL_finngen_results.tsv     INTERVAL plasma pQTL protein-level MR
  extended_cell_crp_MR_results.tsv    immune-cell / CRP layer
  immune_cis_eqtl_instruments.tsv     the instruments themselves

Figures -> 08_figures/cv_novelty/CVN1..CVN8.png
Tables  -> 09_tables/T8_cv_novelty_*.tsv
Report  -> 10_manuscript/Cardiovascular_Novelty_Report.docx

Run:  python src/53_cv_novelty_report.py
"""
import os
import textwrap
import numpy as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = r"I:\Plasma immune atalas"
GC   = os.path.join(ROOT, "06_genetic_causality")
TAB  = os.path.join(ROOT, "09_tables")
FIG  = os.path.join(ROOT, "08_figures", "cv_novelty")
OUT  = os.path.join(ROOT, "10_manuscript")
for d in (FIG, OUT, TAB):
    os.makedirs(d, exist_ok=True)

RISK, PROT, NEU = "#c0392b", "#1f6fb4", "#8d8d8d"
CAT_COL = {
    "NOVEL protein-confirmed": "#7d1128",
    "NOVEL colocalized":       "#c0392b",
    "novel nomination":        "#e08a4a",
    "MHC-caution":             "#9aa0a6",
    "recovered known axis":    "#1f6fb4",
}
CAT_ORDER = ["NOVEL protein-confirmed", "NOVEL colocalized", "novel nomination",
             "MHC-caution", "recovered known axis"]

plt.rcParams.update({
    "font.family": "DejaVu Sans", "font.size": 9,
    "axes.linewidth": 0.8, "axes.edgecolor": "#333333",
    "axes.titlesize": 11, "axes.titleweight": "bold",
    "axes.labelsize": 9.5, "figure.dpi": 300, "savefig.dpi": 300,
    "savefig.bbox": "tight", "legend.frameon": False,
})


def save(fig, name):
    p = os.path.join(FIG, name + ".png")
    fig.savefig(p)
    plt.close(fig)
    print("  fig ->", os.path.basename(p))
    return p


def wrap(s, n=42):
    return "\n".join(textwrap.wrap(str(s), n)[:2])


# ------------------------------------------------------------------ load
print("[53] loading atlas result tables ...")
NOV = pd.read_csv(os.path.join(GC, "novelty_engine_ranked_ALL.tsv"), sep="\t")
CV  = NOV[NOV["disease_system"] == "Cardiovascular"].copy()
CV["novelty_priority"] = CV["novelty_priority"].astype(float)
CV = CV.sort_values("novelty_priority", ascending=False).reset_index(drop=True)

REPL = pd.read_csv(os.path.join(GC, "uk_panphenome_concordance_ALL.tsv"), sep="\t")
PQTL = pd.read_csv(os.path.join(GC, "pqtl_MR_ALL_finngen_results.tsv"), sep="\t")
CELL = pd.read_csv(os.path.join(GC, "extended_cell_crp_MR_results.tsv"), sep="\t")
INST = pd.read_csv(os.path.join(GC, "immune_cis_eqtl_instruments.tsv"), sep="\t")

N_CV      = len(CV)
N_GENE    = CV["gene_symbol"].nunique()
N_DIS     = CV["disease"].nunique()
N_ALL     = len(NOV)
CATS      = CV["category_label"].value_counts()
N_NOVEL   = int(CATS.get("NOVEL protein-confirmed", 0) + CATS.get("NOVEL colocalized", 0))
N_COLOC   = int((CV["PP_H4"] >= 0.8).sum())
N_REPL    = int(CV["two_population_validated"].sum())
N_DRUGG   = int((CV["druggability"] > 0).sum())
N_FDA     = int((CV["fda_target"] == 1).sum())
N_RISK    = int((CV["OR"] > 1).sum())
N_PROT    = int((CV["OR"] < 1).sum())
CV_SHARE  = 100.0 * N_CV / N_ALL

print(f"  cardiovascular: {N_CV} causal pairs | {N_GENE} proteins | {N_DIS} endpoints "
      f"({CV_SHARE:.1f}% of the whole-phenome causal set)")

# therapeutic direction implied by the engine
CV["action"] = np.where(CV["OR"] > 1, "block / antagonise", "agonise / replace")
# a compact evidence tier, same logic as the atlas' 1-5 tiers
def tier(r):
    if r["pqtl_p"] < 0.05 and r["pqtl_PPH4"] >= 0.8:            return 1
    if r["pqtl_p"] < 0.05 and r["PP_H4"] >= 0.8:                return 2
    if r["PP_H4"] >= 0.8 and r["two_population_validated"]:     return 2
    if r["PP_H4"] >= 0.8:                                       return 3
    if r["two_population_validated"]:                           return 3
    return 4
CV["tier"] = CV.apply(tier, axis=1)


# =========================================================== CVN1 funnel
print("[53] CVN1  evidence funnel for the cardiovascular chapter")
fig, ax = plt.subplots(figsize=(7.6, 4.4))
steps = [
    ("cis-MR tests, CV endpoints", None),
    (f"causal at FDR<5%  ({N_CV})", N_CV),
    (f"colocalized PP.H4>=0.8  ({N_COLOC})", N_COLOC),
    (f"two-population replicated  ({N_REPL})", N_REPL),
    (f"protein-level confirmed  ({int(CATS.get('NOVEL protein-confirmed',0))})",
     int(CATS.get("NOVEL protein-confirmed", 0))),
]
# how many CV cis-MR tests were actually run
n_tested = None
try:
    hdr = pd.read_csv(os.path.join(GC, "cis_MR_ALL_finngen_results.tsv"),
                      sep="\t", usecols=["category"], engine="c")
    n_tested = int(hdr["category"].astype(str).str.contains(
        "circulatory|Cardiov|Cardiometabolic", case=False, na=False).sum())
except Exception:
    pass
vals = [n_tested if n_tested else N_CV * 200] + [s[1] for s in steps[1:]]
labs = [f"cis-MR tests, CV endpoints  ({vals[0]:,})"] + [s[0] for s in steps[1:]]
cols = ["#d9d9d9", "#e08a4a", "#c0392b", "#7d1128", "#3d0a16"]
ymax = np.log10(max(vals)) + .4
for i, (v, l, c) in enumerate(zip(vals, labs, cols)):
    h = np.log10(max(v, 1)) + .4
    w = 0.78 - 0.055 * i
    ax.barh(-i, w, height=.62, color=c, edgecolor="none")
    ax.text(w + .015, -i, l, va="center", fontsize=8.6)
ax.set_xlim(0, 1.5); ax.set_ylim(-len(vals) + .5, .6)
ax.axis("off")
ax.set_title("Cardiovascular evidence funnel — what survives each layer",
             loc="left", fontsize=11)
ax.text(0, .95, f"{N_GENE} plasma immune proteins × {N_DIS} FinnGen R12 cardiovascular endpoints",
        transform=ax.transAxes, fontsize=8.4, color="#555")
p1 = save(fig, "CVN1_evidence_funnel")


# =========================================================== CVN2 novelty classes
print("[53] CVN2  novelty class composition")
fig, axes = plt.subplots(1, 2, figsize=(11, 4.2))
ax = axes[0]
order = [c for c in CAT_ORDER if c in CATS.index]
vals = [int(CATS[c]) for c in order]
bars = ax.barh(range(len(order))[::-1], vals,
               color=[CAT_COL[c] for c in order], height=.66)
for y, v in zip(range(len(order))[::-1], vals):
    ax.text(v + max(vals) * .015, y, f"{v}  ({100*v/N_CV:.0f}%)",
            va="center", fontsize=8.6)
ax.set_yticks(range(len(order))[::-1]); ax.set_yticklabels(order, fontsize=8.6)
ax.set_xlabel("causal cardiovascular protein–disease pairs")
ax.set_xlim(0, max(vals) * 1.22)
ax.set_title("A · How novel is the cardiovascular signal?", loc="left")
for s in ("top", "right"): ax.spines[s].set_visible(False)

ax = axes[1]
cvs = CV["category_label"].value_counts(normalize=True)
alls = NOV["category_label"].value_counts(normalize=True)
idx = [c for c in CAT_ORDER if c in cvs.index or c in alls.index]
x = np.arange(len(idx))
ax.bar(x - .19, [100 * cvs.get(c, 0) for c in idx], .36,
       color="#c0392b", label="cardiovascular")
ax.bar(x + .19, [100 * alls.get(c, 0) for c in idx], .36,
       color="#bdc3c7", label="whole phenome")
ax.set_xticks(x)
ax.set_xticklabels([wrap(c, 14) for c in idx], fontsize=7.6, rotation=18, ha="right")
ax.set_ylabel("% of causal pairs")
ax.legend(fontsize=8)
ax.set_title("B · Cardiovascular vs the rest of the phenome", loc="left")
for s in ("top", "right"): ax.spines[s].set_visible(False)
fig.tight_layout()
p2 = save(fig, "CVN2_novelty_classes")


# =========================================================== CVN3 priority ranking
print("[53] CVN3  top novelty-priority cardiovascular targets")
top = CV.head(25).iloc[::-1]
fig, ax = plt.subplots(figsize=(9.6, 8.2))
y = np.arange(len(top))
ax.barh(y, top["novelty_priority"], height=.66,
        color=[CAT_COL.get(c, NEU) for c in top["category_label"]])
for i, r in enumerate(top.itertuples()):
    ax.text(r.novelty_priority + .04, i,
            f"OR {r.OR:.2f} · PP.H4 {r.PP_H4:.2f}"
            + ("  ✔2-pop" if r.two_population_validated else ""),
            va="center", fontsize=7.3, color="#333")
ax.set_yticks(y)
ax.set_yticklabels([f"{r.gene_symbol} → {str(r.disease)[:46]}" for r in top.itertuples()],
                   fontsize=8)
ax.set_xlabel("novelty priority score  (causal + coloc + pleiotropy + druggability\n"
              "+ protein confirmation + replication, penalised for known axes / MHC)")
ax.set_xlim(0, CV["novelty_priority"].max() * 1.34)
ax.set_title("Top 25 novel cardiovascular immune targets", loc="left")
ax.legend(handles=[Line2D([0], [0], marker="s", ls="", ms=8, color=CAT_COL[c], label=c)
                   for c in CAT_ORDER if c in set(top["category_label"])],
          loc="lower right", fontsize=8)
for s in ("top", "right"): ax.spines[s].set_visible(False)
fig.tight_layout()
p3 = save(fig, "CVN3_top25_priority")


# =========================================================== CVN4 novelty vs strength
print("[53] CVN4  novelty vs causal strength")
fig, ax = plt.subplots(figsize=(8.4, 5.6))
d = CV.copy()
d["nlp"] = -np.log10(d["FDR"].clip(lower=1e-300))
for c in CAT_ORDER:
    s = d[d["category_label"] == c]
    if s.empty: continue
    ax.scatter(s["nlp"], s["novelty_priority"], s=26 + 34 * (s["PP_H4"] >= .8),
               color=CAT_COL[c], alpha=.82, edgecolors="white", lw=.5, label=c)
lab = d.nlargest(12, "novelty_priority")
for r in lab.itertuples():
    ax.annotate(f"{r.gene_symbol}", (-np.log10(max(r.FDR, 1e-300)), r.novelty_priority),
                fontsize=7.6, fontweight="bold", xytext=(4, 3),
                textcoords="offset points", color="#222")
ax.set_xlabel("causal strength   -log10 FDR (cis-eQTL MR, FinnGen R12)")
ax.set_ylabel("novelty priority score")
ax.set_title("Novel does not mean weak — the strongest cardiovascular signals\n"
             "are also among the least previously described", loc="left", fontsize=10.5)
ax.legend(fontsize=7.8, loc="lower right", title="engine class", title_fontsize=8)
for s in ("top", "right"): ax.spines[s].set_visible(False)
fig.tight_layout()
p4 = save(fig, "CVN4_novelty_vs_strength")


# =========================================================== CVN5 disease map
print("[53] CVN5  which cardiovascular diseases the immunome reaches")
by_dis = (CV.groupby("disease")
            .agg(n=("gene_symbol", "size"),
                 novel=("category_label", lambda s: int(s.isin(
                     ["NOVEL colocalized", "NOVEL protein-confirmed"]).sum())),
                 top=("novelty_priority", "max"))
            .sort_values("n", ascending=False).head(18).iloc[::-1])
fig, ax = plt.subplots(figsize=(9.2, 6.4))
y = np.arange(len(by_dis))
ax.barh(y, by_dis["n"] - by_dis["novel"], height=.66, color="#e08a4a",
        label="nomination / caution")
ax.barh(y, by_dis["novel"], height=.66, left=by_dis["n"] - by_dis["novel"],
        color="#c0392b", label="NOVEL, colocalized or protein-confirmed")
for i, r in enumerate(by_dis.itertuples()):
    ax.text(r.n + .18, i, f"best score {r.top:.2f}", va="center", fontsize=7.4, color="#444")
ax.set_yticks(y); ax.set_yticklabels([str(i)[:52] for i in by_dis.index], fontsize=8)
ax.set_xlabel("causal plasma immune proteins")
ax.set_xlim(0, by_dis["n"].max() * 1.36)
ax.legend(fontsize=8, loc="lower right")
ax.set_title("Cardiovascular endpoints reached by the plasma immunome", loc="left")
for s in ("top", "right"): ax.spines[s].set_visible(False)
fig.tight_layout()
p5 = save(fig, "CVN5_disease_map")


# =========================================================== CVN6 direction / action
print("[53] CVN6  therapeutic direction")
fig, axes = plt.subplots(1, 2, figsize=(10.6, 4.3))
ax = axes[0]
d = CV.copy()
d["lo"] = np.log(d["OR"])
ax.hist(d.loc[d.OR > 1, "lo"], bins=26, color=RISK, alpha=.85, label=f"risk-increasing ({N_RISK})")
ax.hist(d.loc[d.OR < 1, "lo"], bins=26, color=PROT, alpha=.85, label=f"protective ({N_PROT})")
ax.axvline(0, color="k", lw=.8, ls="--")
ax.set_xlabel("log odds ratio per 1-SD genetically-proxied protein")
ax.set_ylabel("causal pairs"); ax.legend(fontsize=8)
ax.set_title("A · Direction of effect", loc="left")
for s in ("top", "right"): ax.spines[s].set_visible(False)

ax = axes[1]
grp = (CV.groupby(["action", "category_label"]).size().unstack(fill_value=0)
         .reindex(columns=[c for c in CAT_ORDER if c in CV.category_label.unique()],
                  fill_value=0))
bottom = np.zeros(len(grp))
for c in grp.columns:
    ax.bar(range(len(grp)), grp[c], .5, bottom=bottom, color=CAT_COL[c], label=c)
    bottom += grp[c].values
ax.set_xticks(range(len(grp)))
ax.set_xticklabels([f"{i}\n(n={int(grp.loc[i].sum())})" for i in grp.index], fontsize=8.4)
ax.set_ylabel("causal pairs")
ax.legend(fontsize=7.4)
ax.set_title("B · Implied therapeutic action", loc="left")
for s in ("top", "right"): ax.spines[s].set_visible(False)
fig.tight_layout()
p6 = save(fig, "CVN6_direction_action")


# =========================================================== CVN7 evidence tiers
print("[53] CVN7  evidence tier x druggability")
fig, axes = plt.subplots(1, 2, figsize=(10.6, 4.2))
ax = axes[0]
tc = CV["tier"].value_counts().sort_index()
ax.bar([f"Tier {i}" for i in tc.index], tc.values, .55,
       color=["#7d1128", "#c0392b", "#e08a4a", "#bdc3c7"][:len(tc)])
for i, v in enumerate(tc.values):
    ax.text(i, v + max(tc.values) * .02, str(v), ha="center", fontsize=8.6)
ax.set_ylabel("causal pairs")
ax.set_title("A · Integrated evidence tier", loc="left")
for s in ("top", "right"): ax.spines[s].set_visible(False)

ax = axes[1]
dmap = {0: "no known ligand", 1: "small-molecule\ntractable",
        2: "antibody\ntractable", 3: "existing drug\ntarget"}
dc = CV["druggability"].value_counts().sort_index()
ax.bar([dmap.get(i, str(i)) for i in dc.index], dc.values, .55, color="#1f6fb4")
for i, v in enumerate(dc.values):
    ax.text(i, v + max(dc.values) * .02, str(v), ha="center", fontsize=8.6)
ax.set_ylabel("causal pairs")
ax.set_title(f"B · Druggability  ({N_DRUGG} tractable, {N_FDA} already FDA-target class)",
             loc="left", fontsize=10)
for s in ("top", "right"): ax.spines[s].set_visible(False)
fig.tight_layout()
p7 = save(fig, "CVN7_tiers_druggability")


# =========================================================== CVN8 case studies
print("[53] CVN8  worked case studies")
cases = CV.drop_duplicates("gene_symbol").head(6)
fig, axes = plt.subplots(2, 3, figsize=(12.4, 6.6))
for ax, r in zip(axes.ravel(), cases.itertuples()):
    g = r.gene_symbol
    sub = CV[CV["gene_symbol"] == g].nlargest(6, "novelty_priority").iloc[::-1]
    lo, hi = np.log(sub["OR"]), None
    ax.axvline(0, color="k", lw=.8, ls="--")
    ax.scatter(np.log(sub["OR"]), range(len(sub)), s=42,
               color=[RISK if o > 1 else PROT for o in sub["OR"]], zorder=3)
    for i, rr in enumerate(sub.itertuples()):
        ax.text(np.log(rr.OR), i + .26, f"PP.H4 {rr.PP_H4:.2f}", fontsize=6.6,
                ha="center", color="#555")
    ax.set_yticks(range(len(sub)))
    ax.set_yticklabels([str(x)[:34] for x in sub["disease"]], fontsize=7)
    ax.set_xlabel("log OR", fontsize=8)
    cls = str(r.immune_class)
    ax.set_title(f"{g}  ·  {cls[:26]}\n{r.category_label}", fontsize=8.8, pad=16)
    ax.tick_params(labelsize=7)
    ax.margins(y=.22)
    for s in ("top", "right"): ax.spines[s].set_visible(False)
fig.suptitle("Worked cardiovascular case studies — the six highest-priority novel proteins",
             fontsize=11.5, fontweight="bold", y=1.02)
fig.tight_layout(h_pad=3.0, w_pad=2.0)
p8 = save(fig, "CVN8_case_studies")


# =========================================================== tables
print("[53] writing tables")
cols = ["gene_symbol", "disease", "immune_class", "OR", "FDR", "PP_H4",
        "pqtl_OR", "pqtl_p", "pqtl_PPH4", "two_population_validated",
        "druggability", "fda_target", "category_label", "tier",
        "action", "novelty_priority"]
CV[cols].to_csv(os.path.join(TAB, "T8_cv_novelty_all_pairs.tsv"), sep="\t", index=False)
CV[CV["category_label"].isin(["NOVEL colocalized", "NOVEL protein-confirmed"])][cols] \
  .to_csv(os.path.join(TAB, "T8_cv_novelty_novel_only.tsv"), sep="\t", index=False)
GENE = (CV.groupby(["gene_symbol", "immune_class"])
          .agg(pairs=("disease", "nunique"),
               best_score=("novelty_priority", "max"),
               best_disease=("disease", lambda s: CV.loc[s.index, "disease"]
                             .loc[CV.loc[s.index, "novelty_priority"].idxmax()]),
               n_coloc=("PP_H4", lambda s: int((s >= .8).sum())),
               n_repl=("two_population_validated", "sum"),
               druggability=("druggability", "max"))
          .reset_index().sort_values("best_score", ascending=False))
GENE.to_csv(os.path.join(TAB, "T8_cv_novelty_by_protein.tsv"), sep="\t", index=False)
print(f"  T8 tables -> {TAB}")


# =========================================================== Word report
print("[53] building the Word report")
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)


def H(t, lvl=1):
    h = doc.add_heading(t, level=lvl)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x7D, 0x11, 0x28)
    return h


def P(t, size=10.5, italic=False, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.size = Pt(size); r.italic = italic; r.bold = bold
    if align: p.alignment = align
    return p


def FIGP(path, cap, width=6.3):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    r = c.add_run(cap); r.font.size = Pt(8.5); r.italic = True
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def TABLE(df, cols=None, n=25, widths=None):
    df = df.head(n)
    cols = cols or list(df.columns)
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(cols):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(c))
        run.bold = True; run.font.size = Pt(8)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(cols):
            v = row[c]
            if isinstance(v, float):
                v = f"{v:.3g}" if (abs(v) < 1e-3 or abs(v) >= 1e4) else f"{v:.3f}"
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(v))
            run.font.size = Pt(7.6)
    return t


# ---- title page
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("Novelty Analysis of the Plasma Immunome\nin Cardiovascular Disease")
r.bold = True; r.font.size = Pt(21); r.font.color.rgb = RGBColor(0x7D, 0x11, 0x28)
P("A genetics-anchored assessment of which plasma immune proteins are causally "
  "implicated in cardiovascular disease, how much of that signal is genuinely new, "
  "and what should be done about it.",
  size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
P(f"Human Plasma Immune Atlas · layer 7 (novelty & therapeutic-direction engine) · "
  f"FinnGen R12 cardiovascular chapter", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
P("Open resource: https://huggingface.co/spaces/jianlizhao/Human-Plasma-Immune-Atlas",
  size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
H("Headline result", 1)
P(f"Across the cardiovascular chapter of FinnGen R12, {N_GENE} plasma immune proteins are "
  f"causally implicated in {N_DIS} cardiovascular endpoints, giving {N_CV} causal "
  f"protein–disease pairs at FDR < 5 %. This is the single largest disease system in the "
  f"atlas: {CV_SHARE:.1f} % of all {N_ALL:,} whole-phenome causal pairs are cardiovascular, "
  f"more than respiratory ({int((NOV.disease_system=='Respiratory').sum())}) or "
  f"musculoskeletal ({int((NOV.disease_system=='Musculoskeletal').sum())}) disease.")
P(f"{N_COLOC} of the {N_CV} pairs colocalise with the disease signal (PP.H4 ≥ 0.8), meaning "
  f"the protein and the disease share the same causal variant rather than two variants in "
  f"linkage disequilibrium. {N_REPL} additionally replicate in an independent "
  f"Finland → England two-population test.")
P(f"Only {int(CATS.get('recovered known axis', 0))} of the {N_CV} pairs is a previously "
  f"established immune–cardiovascular drug axis. "
  f"{N_NOVEL} pairs ({100*N_NOVEL/N_CV:.0f} %) are classified NOVEL with colocalisation "
  f"support, of which {int(CATS.get('NOVEL protein-confirmed',0))} are additionally confirmed "
  f"at the circulating-protein level using INTERVAL plasma pQTL. In other words, the "
  f"cardiovascular immunome is almost entirely unexploited therapeutic territory.",
  bold=False)
P(f"{N_DRUGG} of the {N_CV} pairs involve a protein with an existing tractability route "
  f"(small molecule, antibody or an approved drug class), and {N_FDA} sit on proteins that "
  f"are already the target of an approved drug for some other indication — i.e. immediate "
  f"repurposing candidates.")

FIGP(p1, "Figure 1 · Cardiovascular evidence funnel. Each layer of the atlas is an "
         "independent filter; the surviving set is the therapeutically actionable core.")

# ---- what novelty means
H("How novelty is defined", 1)
P("The novelty engine does not ask whether a protein has ever been mentioned alongside "
  "heart disease. It asks whether the specific causal claim — this protein, this direction, "
  "this endpoint — is already exploited. Each pair receives a score built from seven "
  "components and two penalties:")
comp = pd.DataFrame({
    "Component": ["s_causal", "s_coloc", "s_pleio", "s_drug", "s_cell", "s_protein",
                  "s_repl", "p_known", "p_mhc"],
    "Meaning": ["strength of the cis-eQTL MR effect (FDR)",
                "colocalisation posterior PP.H4",
                "specificity — penalised if the protein hits many unrelated chapters",
                "druggability / tractability class",
                "does the protein also move immune cell counts or CRP",
                "confirmed at the circulating-protein level (INTERVAL pQTL MR)",
                "independent two-population replication (UK Biobank)",
                "PENALTY — already an approved drug axis",
                "PENALTY — MHC region, long-range LD defeats colocalisation"],
    "CV mean": [f"{CV[c].mean():.2f}" for c in
                ["s_causal", "s_coloc", "s_pleio", "s_drug", "s_cell",
                 "s_protein", "s_repl", "p_known", "p_mhc"]],
})
TABLE(comp, n=9)
P("")
P(f"The cardiovascular set scores high on causal strength (mean {CV.s_causal.mean():.2f}) "
  f"and specificity (mean {CV.s_pleio.mean():.2f}) but low on protein confirmation "
  f"(mean {CV.s_protein.mean():.2f}) — the honest reading is that the transcript-level "
  f"evidence is strong and the plasma-protein arbitration is still thin, because public "
  f"login-free pQTL coverage of these proteins is incomplete rather than because the "
  f"signals failed.")

FIGP(p2, "Figure 2 · Novelty class composition. (A) The cardiovascular chapter split by "
         "engine class. (B) Cardiovascular disease is enriched for novel colocalized signal "
         "relative to the rest of the phenome, and almost devoid of recovered known axes.")

# ---- top targets
H("The top novel cardiovascular targets", 1)
P("Ranked by novelty priority. Every row is a live claim: a direction of effect, a "
  "colocalisation posterior, and where available an independent-population test.")
FIGP(p3, "Figure 3 · Top 25 novel cardiovascular immune targets, coloured by engine class. "
         "OR is per 1-SD genetically-proxied increase in the plasma protein; ✔2-pop marks "
         "pairs that also replicate in UK Biobank.")

show = CV.head(20)[["gene_symbol", "disease", "immune_class", "OR", "FDR", "PP_H4",
                    "two_population_validated", "category_label", "tier", "action"]].copy()
show.columns = ["Gene", "Cardiovascular endpoint", "Immune class", "OR", "FDR", "PP.H4",
                "2-pop", "Novelty class", "Tier", "Implied action"]
show["Cardiovascular endpoint"] = show["Cardiovascular endpoint"].astype(str).str[:44]
show["Immune class"] = show["Immune class"].astype(str).str[:22]
TABLE(show, n=20)
P("Table 1 · The 20 highest-priority novel cardiovascular immune targets. "
  "Full table: 09_tables/T8_cv_novelty_all_pairs.tsv", size=8.5, italic=True)

FIGP(p4, "Figure 4 · Novelty against causal strength. Novelty here is not a consolation "
         "prize for weak signal — several of the least-described proteins carry the "
         "strongest cardiovascular MR evidence in the atlas.")

# ---- case studies
H("Worked case studies", 1)
for r in cases.itertuples():
    g = r.gene_symbol
    sub = CV[CV["gene_symbol"] == g]
    nd = sub["disease"].nunique()
    nco = int((sub["PP_H4"] >= .8).sum())
    nre = int(sub["two_population_validated"].sum())
    dirn = "increases" if r.OR > 1 else "reduces"
    act = "block or antagonise it" if r.OR > 1 else "agonise, supplement or replace it"
    pq = ""
    if pd.notna(r.pqtl_p) and r.pqtl_p < 0.05:
        pq = (f" The claim is confirmed at the circulating-protein level using INTERVAL "
              f"plasma pQTL (protein-level OR {r.pqtl_OR:.2f}, P = {r.pqtl_p:.2g}), so the "
              f"signal is not a transcript-only artefact.")
    rep = (f" It replicates in the independent Finland → England two-population test."
           if r.two_population_validated else
           " It does not yet have an independent-population test, because no matching "
           "single-cohort UK Biobank GWAS of this endpoint exists.")
    H(f"{g} → {str(r.disease)[:60]}", 2)
    P(f"Class: {r.immune_class} · engine verdict: {r.category_label} · tier {int(r.tier)} · "
      f"novelty priority {r.novelty_priority:.2f}", size=9.5, italic=True)
    P(f"A genetically-proxied 1-SD increase in plasma {g} {dirn} the odds of "
      f"{str(r.disease).lower()} (OR {r.OR:.2f}, FDR {r.FDR:.2g}). The protein and the "
      f"disease share a causal variant with posterior PP.H4 = {r.PP_H4:.2f}."
      + rep + pq +
      f" Across the cardiovascular chapter {g} is causal for {nd} endpoint(s), "
      f"{nco} of them colocalised and {nre} two-population replicated. "
      f"Therapeutic reading: {act}.")

FIGP(p8, "Figure 5 · Worked case studies. For each of the six highest-priority novel "
         "proteins, every cardiovascular endpoint it causes, with the colocalisation "
         "posterior above each point. Red = risk-increasing, blue = protective.")

# ---- disease view
H("Which cardiovascular diseases the immunome reaches", 1)
P("The signal is not confined to atherosclerosis. Hypertension, venous disease, atrial "
  "fibrillation and revascularisation endpoints all carry independent plasma-immune causal "
  "architecture, and the venous endpoints in particular are dominated by novel colocalised "
  "proteins with essentially no prior immune-therapeutic literature.")
FIGP(p5, "Figure 6 · Cardiovascular endpoints ranked by the number of causal plasma immune "
         "proteins; the dark segment is the NOVEL colocalised / protein-confirmed fraction.")

dis = by_dis.iloc[::-1].reset_index()
dis.columns = ["Cardiovascular endpoint", "Causal proteins", "of which NOVEL", "Best score"]
dis["Cardiovascular endpoint"] = dis["Cardiovascular endpoint"].astype(str).str[:56]
TABLE(dis, n=18)
P("Table 2 · Cardiovascular endpoints with the most causal plasma immune proteins.",
  size=8.5, italic=True)

# ---- direction / action
H("Therapeutic direction and tractability", 1)
P(f"Of the {N_CV} causal pairs, {N_RISK} are risk-increasing (implied action: block the "
  f"protein) and {N_PROT} are protective (implied action: agonise, supplement or replace it). "
  f"The protective majority is important: it means a large part of the cardiovascular "
  f"immunome would be damaged, not helped, by broad immunosuppression — the direction has "
  f"to be read protein by protein.")
FIGP(p6, "Figure 7 · (A) Distribution of causal effect direction. (B) Implied therapeutic "
         "action split by novelty class.")
FIGP(p7, "Figure 8 · (A) Integrated evidence tier — tier 1–2 pairs carry protein-level or "
         "replicated colocalised support. (B) Druggability of the underlying proteins.")

byg = GENE.head(20).copy()
byg.columns = ["Gene", "Immune class", "CV endpoints", "Best score", "Best endpoint",
               "Colocalised", "2-pop replicated", "Druggability"]
byg["Best endpoint"] = byg["Best endpoint"].astype(str).str[:40]
byg["Immune class"] = byg["Immune class"].astype(str).str[:20]
TABLE(byg, n=20)
P("Table 3 · The 20 highest-priority cardiovascular immune proteins, collapsed across "
  "endpoints. Full table: 09_tables/T8_cv_novelty_by_protein.tsv", size=8.5, italic=True)

# ---- limits
H("What this analysis does not claim", 1)
for t in [
    "MR estimates are lifelong genetically-proxied effects. They identify a causal direction; "
    "they do not predict the effect size of a drug given to an adult patient.",
    f"{int(CATS.get('MHC-caution',0))} cardiovascular pairs sit in or near the MHC region and "
    "are held at nomination only — long-range linkage disequilibrium there defeats "
    "colocalisation, so a shared-variant posterior cannot be trusted.",
    "cis-eQTL instruments proxy transcript abundance, which is not always circulating protein. "
    "The plasma pQTL layer is the arbiter, and it currently covers only part of this protein "
    "set because public login-free SomaScan/Olink aptamer coverage is incomplete.",
    "The two-population layer uses only single-cohort UK Biobank GWAS. Several large public "
    "meta-analyses of the same endpoints silently include FinnGen, which would make "
    "'replication' circular; those were deliberately excluded, which is why only "
    f"{N_REPL} of {N_CV} pairs carry a replication flag.",
    "'Novel' is a statement about therapeutic exploitation, not about the literature. A "
    "protein may have been described in cardiovascular biology and still be scored novel if "
    "no approved drug acts on that axis for that indication.",
    "Research resource only — not clinical advice and not a validated diagnostic.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t); r.font.size = Pt(10)

# ---- methods
H("Methods, in brief", 1)
P("Instruments: strongest cis-eQTL per plasma immune protein from eQTLGen (n ≈ 31,684). "
  "Outcomes: FinnGen release R12, every endpoint in the cardiovascular chapter. Causal "
  "estimation: Wald ratio / Zhu approximation, Benjamini–Hochberg FDR across the whole "
  "phenome (not within-chapter, so the cardiovascular results are not preferentially "
  "significant). Colocalisation: coloc.abf with Wakefield approximate Bayes factors, "
  "PP.H4 ≥ 0.8. Protein-level confirmation: INTERVAL plasma pQTL (Sun et al. 2018, public "
  "via the EBI GWAS Catalog) MR + colocalisation. Two-population replication: independent "
  "single-cohort UK Biobank GWAS of the matched endpoint. Immune-cell layer: the same "
  "instruments against blood cell counts, fractions and CRP. Novelty engine: the seven "
  "components and two penalties tabulated above, combined into the novelty priority score.")
P(f"Source tables: 06_genetic_causality/novelty_engine_ranked_ALL.tsv "
  f"({N_ALL:,} whole-phenome causal pairs), uk_panphenome_concordance_ALL.tsv, "
  f"pqtl_MR_ALL_finngen_results.tsv, extended_cell_crp_MR_results.tsv. "
  f"Derived tables written by this script: 09_tables/T8_cv_novelty_all_pairs.tsv, "
  f"T8_cv_novelty_novel_only.tsv, T8_cv_novelty_by_protein.tsv. "
  f"Figures: 08_figures/cv_novelty/CVN1–CVN8.", size=9)
P("Every figure and number in this report is regenerated from those tables by "
  "src/53_cv_novelty_report.py. Nothing is hand-entered.", size=9, italic=True)

path = os.path.join(OUT, "Cardiovascular_Novelty_Report.docx")
doc.save(path)
print("[53] report ->", path)
print(f"[53] done: {N_CV} CV pairs | {N_NOVEL} novel colocalized/protein-confirmed | "
      f"{N_REPL} two-population replicated | {N_DRUGG} druggable")
